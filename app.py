import streamlit as st
import google.generativeai as genai
import io
import base64
import json
import time
import os
from datetime import datetime
import fitz  # PyMuPDF for PDF text extraction
import re
import pandas as pd
import matplotlib.pyplot as plt
import plotly.express as px
import seaborn as sns
import pyperclip


# Configure Streamlit page
st.set_page_config(page_title="Ultimate AI Creator Hub", page_icon="🧠", layout="wide", initial_sidebar_state="expanded")

# Custom CSS (highly condensed)
st.markdown("""<style>.main{background-color:#f8f9fa}.stApp{max-width:1200px;margin:0 auto}.tool-category{font-size:1.2rem;font-weight:bold;margin-top:1rem;color:#1e3a8a}div[data-testid="stVerticalBlock"]{gap:0.5rem}.stButton>button{background:linear-gradient(90deg,#3b82f6,#8b5cf6);color:white;font-weight:600;border-radius:10px;padding:10px 20px;box-shadow:0 4px 14px rgba(0,0,0,0.1);transition:all 0.3s ease}.stButton>button:hover{transform:translateY(-2px);box-shadow:0 6px 20px rgba(0,0,0,0.15)}div.stTabs [data-baseweb="tab-list"]{gap:8px}div.stTabs [data-baseweb="tab"]{background-color:#300661;border-radius:8px 8px 0px 0px;padding:10px 16px;font-weight:600}div.stTabs [aria-selected="true"]{background-color:#450542}.output-box{background-color:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;padding:20px;margin-top:20px}.history-item{padding:10px;border:1px solid #e5e7eb;border-radius:8px;margin-bottom:8px;background-color:white}.category-selector{margin:10px 0}.search-box{margin:15px 0}</style>""", unsafe_allow_html=True)

# Initialize session state
if 'api_key' not in st.session_state: st.session_state.api_key = ""
if 'history' not in st.session_state: st.session_state.history = []
if 'api_model' not in st.session_state: st.session_state.api_model = "gemini-2.0-flash"
if 'prompt_templates' not in st.session_state: st.session_state.prompt_templates = {}



import os
import subprocess
import sys

# Function to check & install system dependencies (Tesseract & FFmpeg)
def install_dependencies():
    try:
        # Check if Tesseract is installed
        subprocess.run(["tesseract", "-v"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print("✅ Tesseract is already installed.")
    except FileNotFoundError:
        print("⚠️ Tesseract not found. Installing now...")
        if os.name == "nt":  # Windows
            os.system("winget install -e --id UB-Mannheim.TesseractOCR")  # Requires Windows Package Manager
        elif os.name == "posix":  # Linux & MacOS
            os.system("sudo apt update && sudo apt install -y tesseract-ocr")
        print("✅ Tesseract installed successfully.")

    try:
        # Check if FFmpeg is installed
        subprocess.run(["ffmpeg", "-version"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print("✅ FFmpeg is already installed.")
    except FileNotFoundError:
        print("⚠️ FFmpeg not found. Installing now...")
        if os.name == "nt":  # Windows
            os.system("winget install -e --id Gyan.FFmpeg")  # Requires Windows Package Manager
        elif os.name == "posix":  # Linux & MacOS
            os.system("sudo apt update && sudo apt install -y ffmpeg")
        print("✅ FFmpeg installed successfully.")

# Run dependency check on startup
install_dependencies()



import pytesseract
from PIL import Image

# Function to extract text from images (OCR)
def extract_text_from_image(uploaded_file):
    """Extracts text from an uploaded image file using OCR."""
    image = Image.open(uploaded_file)
    text = pytesseract.image_to_string(image)
    return text if text.strip() else "No text detected in the image."


# Function to extract text from CSV files
def extract_text_from_csv(uploaded_file):
    """Reads a CSV file and returns its content as a formatted string."""
    df = pd.read_csv(uploaded_file)
    return df.to_string(index=False)  # Convert DataFrame to readable text



# Function to extract text from TXT files
def extract_text_from_txt(uploaded_file):
    """Reads and returns text from a TXT file."""
    return uploaded_file.read().decode("utf-8")


import docx  # Ensure you have python-docx installed

# Function to extract text from DOCX files
def extract_text_from_docx(uploaded_file):
    """Reads and extracts text from a DOCX file."""
    doc = docx.Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])  # Join paragraphs into a single text


import speech_recognition as sr
from pydub import AudioSegment
import io

# Function to extract text from audio files
def extract_text_from_audio(uploaded_file):
    """Converts speech from an audio file (MP3/WAV) to text using Google Speech Recognition."""
    recognizer = sr.Recognizer()
    
    # Read the file into a binary stream
    audio_bytes = uploaded_file.read()

    # Convert MP3 to WAV if needed
    try:
        audio = AudioSegment.from_file(io.BytesIO(audio_bytes))  # Auto-detect format
        audio = audio.set_channels(1).set_frame_rate(16000)  # Convert to mono, 16kHz
        wav_bytes = io.BytesIO()
        audio.export(wav_bytes, format="wav")
        wav_bytes.seek(0)  # Move to start of file
    except Exception as e:
        return f"Error processing audio file: {str(e)}"

    # Transcribe the converted WAV file
    with sr.AudioFile(wav_bytes) as source:
        audio_data = recognizer.record(source)

    try:
        text = recognizer.recognize_google(audio_data)  # Uses Google Speech Recognition API
        return text if text.strip() else "No speech detected in the audio."
    except sr.UnknownValueError:
        return "Could not understand the audio."
    except sr.RequestError:
        return "Could not request results from Google Speech Recognition."

import fitz  # For PDF processing (PyMuPDF)
import pandas as pd
import docx

def extract_text_from_file(uploaded_file):
    """Extracts text from an uploaded file (PDF, DOCX, TXT, CSV)."""
    file_type = uploaded_file.type
    if file_type == "application/pdf":
        # For PDFs
        doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        return "\n".join([page.get_text("text") for page in doc])
    elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        # For DOCX files
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_type in ["text/plain"]:
        # For TXT files
        return uploaded_file.read().decode("utf-8")
    elif file_type in ["text/csv"]:
        # For CSV files
        df = pd.read_csv(uploaded_file)
        return df.to_string(index=False)
    else:
        return "Unsupported file format."


import re

def extract_references(text):
    """Extracts possible reference sections from research papers."""
    references = re.findall(r"\[.*?\]|\(.*?\)", text)
    return references if references else ["No references detected."]


import fitz  # PyMuPDF for PDF text extraction

def extract_text_from_pdf(uploaded_file):
    """Extracts text from an uploaded PDF file."""
    try:
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as doc:
            text = "\n".join([page.get_text("text") for page in doc])
        return text
    except Exception as e:
        return f"Error extracting text: {str(e)}"

# Function to generate expanded AI tools list
@st.cache_data
def generate_ai_tools():
    # Base categories dictionary - we'll expand with multipliers later
    base_categories = {
        "Writing": [
            # Professional & Business Writing  
            "Resume", "Cover Letter", "Job Application Letter", "Recommendation Letter", "Business Proposal",  
            "Project Proposal", "Executive Summary", "Professional Bio", "Company Profile", "Business Plan",  
            "Marketing Copy", "Sales Pitch", "Brand Story", "Press Release", "Meeting Minutes", "Email",  
            "SOP (Standard Operating Procedure)", "RFP (Request for Proposal)", "Policy Draft", "Employee Handbook",  
            "Business Case Study", "Whitepaper",  

            # Academic & Research Writing  
            "Academic Essay", "Thesis Statement", "Literature Review", "Research Paper", "Citation",  
            "Abstract & Keywords Generator", "Lab Report", "Case Study", "Grant Proposal", "Conference Paper",  
            "Technical Report", "Editorial Guidelines", "Style Guide",  

            # Creative & Content Writing  
            "Blog Post", "Script (Movie, YouTube, Podcast)", "Short Story", "Novel Outline", "Poetry",  
            "Creative Writing Prompt", "Character Development", "Plot Outline", "Ghostwriting",  
            "Content Ideation & Brainstorming", "Social Media Caption & Hashtag Generator",  

            # Legal & Compliance Writing  
            "Legal Document Drafting", "Contract Clause Generator", "Terms & Conditions", "Privacy Policy",  
            "GDPR Compliance Notice", "Non-Disclosure Agreement (NDA)", "Service-Level Agreement (SLA)",  
            "Disclaimers & Liability Waivers",  

            # Editing & Optimization  
            "Grammar & Spell Check", "Paraphrasing Tool", "Readability Improvement", "Plagiarism Checker",  
            "Tone Adjustment (Formal, Casual, Professional)", "Style Consistency Checker",  
            "Sentence Expansion & Condensation",  

            # Summarization & Extraction  
            "Executive Summary Generator", "Meeting Notes & Minutes Summarization", "Article Summary",  
            "Research Paper Summary", "Legal Document Summary", "Report Extraction", "Transcript Summarization",  

            # Technical & Documentation Writing  
            "API Documentation Generator", "Software User Manual", "Product Documentation",  
            "Engineering Report", "Patent Writing", "Troubleshooting Guide", "FAQs & Help Desk Content",  

            # UX & Product Writing  
            "UX Writing (Microcopy for Apps & Websites)", "Chatbot Script", "Error Message Optimization",  
            "App Store Description", "Product Descriptions",  

            # Miscellaneous Writing  
            "Personal Letter (Apology, Love, Invitation)", "Speech Writing (Wedding, Graduation, Motivational)",  
            "Eulogy Writing", "Greeting Card Messages", "Review Writing (Product, Movie, Book)"  
        ],

        
        "Creative": [
            "Poem", "Story", "Dialogue", "Character", "Book Title", "Horror Story", "Sci-Fi Story",
            "Song Lyrics", "Children's Story", "Novel Outline", "Metaphor", "Joke", "Fantasy World",
            "Sci-Fi Technology", "Historical Fiction", "Memoir", "Poetry Prompt", "Creative Prompt",
            "Short Story Starter", "Screenplay Format", "Plot Twist", "Character Backstory",
            "Setting Description", "Alternate History", "Mythical Creature", "Magic System",
            "Fictional Language", "Story Conflict","Dream Interpreter", "Symbolic Dream Meaning", "Lucid Dream Story", "Surreal Dream Narration", "Dream Journal Entry", "Visionary Experience"  
        ],

        
        "Business": [
            # Business Strategy & Planning  
            "Business Idea", "Business Model Canvas", "Business Plan", "Go-to-Market Strategy",  
            "Startup Pitch", "Investor Pitch", "Funding Request", "Bootstrapping Strategy",  
            "Growth Hacking Techniques", "Project Timeline", "Risk Assessment", "Exit Strategy",  

            # Marketing & Branding  
            "Marketing Strategy", "Content Marketing Plan", "Branding Guide", "Social Media Strategy",  
            "SEO Keywords", "Ad Copy Generator", "Product Launch Plan", "Rebranding Strategy",  
            "Customer Retention Plan", "Brand Voice & Messaging",  

            # Sales & Customer Development  
            "Sales Pitch", "Freelance Proposal", "Lead Generation Strategy", "Customer Persona",  
            "Cold Email Template", "Sales Funnel Optimization", "Customer Support Guidelines",  
            "Loyalty Program Strategy", "Upselling & Cross-Selling Strategy",  

            # Financial & Business Analysis  
            "Grant Proposal", "ROI Calculator", "Break-even Analysis", "Pricing Strategy",  
            "Financial Forecasting", "Cash Flow Management", "Profit Margin Analysis",  
            "Funding Proposal", "Investor Report",  

            # Market Research & Competitive Analysis  
            "Market Research", "Competitor Analysis", "Industry Trends Report", "Customer Survey Questions",  
            "Business Case", "SWOT Analysis", "KPI Framework", "Benchmarking Analysis",  

            # Business Communication & Operations  
            "Business Email", "LinkedIn Bio", "Professional Networking Message", "Press Release",  
            "Mission Statement", "Company Values", "Code of Conduct", "HR Policy Draft",  
            "Crisis Management Plan", "Standard Operating Procedure (SOP)",  

            # Entrepreneurship & Consultation  
            "Business Consultation", "Startup Incubation Strategy", "Fundraising Strategy",  
            "Legal Compliance Checklist", "Franchise Business Plan", "Freelance Business Plan",  
            "E-commerce Strategy", "Subscription Model Strategy"  
        ]

    }

    
    # Additional categories to reach 2000+ tools
    extended_categories = {
        "Social Media": [
            # General Social Media Content  
            "Post", "Caption", "Viral Tweet", "Hashtag Strategy", "Content Calendar", "Trending Topics Finder",  
            "Community Post", "Meme Generator", "User Engagement Prompt", "Social Media Poll Idea",  

            # YouTube & Video Content  
            "YouTube Idea", "YouTube Script", "YouTube Thumbnail Text", "Reel Script",  
            "TikTok Trend", "Instagram Story", "Short-Form Video Hook", "Livestream Outline",  

            # Influencer & Brand Marketing  
            "Influencer Pitch", "Brand Partnership", "Sponsorship Email", "Affiliate Marketing Copy",  
            "Networking Opener", "Collaborative Post Idea", "Product Placement Script",  

            # Advertising & Growth Strategies  
            "Facebook Ad", "Instagram Ad Copy", "LinkedIn Ad Copy", "Twitter Ad Copy",  
            "A/B Testing Ad Variations", "Social Media Growth Hack", "Viral Formula",  

            # Engagement & Audience Interaction  
            "Review Response", "Crisis Response", "Customer Support Reply", "Community Engagement Strategy",  
            "DM Outreach Message", "Story Poll & Quiz Ideas",  

            # Niche-Specific Content  
            "Podcast Episode Idea", "Podcast Episode Title", "Podcast Teaser Post", "Dating Profile",  
            "LinkedIn Article", "Twitter Thread", "Pinterest Description", "Tagline",  
            "Event Promotion Post", "Giveaway & Contest Post",  

            # SEO & Algorithm Boosting Content  
            "SEO-Optimized Post", "Trending Hashtag Suggestions", "Social Media Headline Generator",  
            "Evergreen Content Idea", "Cross-Platform Content Repurposing"  
        ],
        
        "Productivity": [
            # Planning & Scheduling  
            "Productivity Plan", "Daily Plan", "Weekly Schedule", "Monthly Planner", "Yearly Goal Planner",  
            "Project Management", "Task Breakdown", "Meeting Agenda", "Time Blocking Planner",  
            "Work-Life Balance Plan", "Pomodoro Session Planner",  

            # Organization & Efficiency  
            "Note-Taking", "Brainstorming", "Mind Mapping", "Decision Matrix", "Prioritization",  
            "Time Management", "Kanban Board Organizer", "Checklist Generator", "Routine Optimization",  

            # Goal Setting & Self-Improvement  
            "Goal Setting", "Habit Tracker", "Professional Development Plan", "Self-Reflection Journal",  
            "Personal SWOT Analysis", "Skill Development Roadmap", "Morning & Night Routine Guide",  

            # Travel & Logistics  
            "Travel Itinerary", "Packing List", "Grocery List", "Meal Planning", "Event Planning",  

            # Career & Productivity Tools  
            "Interview Prep", "Performance Review Template", "Networking Plan", "Work Presentation Structure",  
            "Work Journal", "Bullet Journal Prompts", "Crisis Management Plan"  
        ],
        
        "Education": [
            # Teaching & Curriculum Design  
            "Lesson Plan", "Course Outline", "Curriculum", "Syllabus", "Workshop Plan",  
            "Classroom Activity Idea", "Lecture Notes", "Tutoring Session Plan",  

            # Student Learning & Study Aids  
            "Educational Quiz", "Study Guide", "Subject Summary", "Assignment", "Test Questions",  
            "Flashcard Generator", "Formula Cheat Sheet", "Learning Path", "Research Topic Suggestions",  

            # Interactive & Alternative Learning  
            "Educational Game", "Puzzle-Based Learning", "Simulation-Based Teaching", "Case Study Analysis",  
            "Peer Learning Strategy", "Role-Playing Educational Exercise",  

            # Higher Education & Academic Tools  
            "Academic Resource List", "Research Paper Guide", "Thesis Outline", "Citation Generator",  
            "Critical Thinking Exercise", "Problem-Solving Challenge",  

            # Skill Development & Self-Learning  
            "Tutorial", "Self-Study Plan", "Lifelong Learning Roadmap", "Language Learning Plan",  
            "Career-Oriented Learning Path", "Skill-Building Guide"  
        ],
        
        "Design": [
            # Branding & Identity  
            "Design Brief", "Logo Concept", "Brand Identity", "Typography Guide", "Color Palette",  
            "Mood Board", "Style Tile", "Visual Hierarchy Guide",  

            # UI/UX & Digital Design  
            "UI Element", "UX Flow", "Wireframe Sketch", "Website Layout", "Mobile App UI",  
            "Dashboard UI", "Landing Page Wireframe", "Accessibility Guidelines",  

            # Print & Packaging Design  
            "Print Material", "Product Packaging", "Brochure Layout", "Business Card Design",  
            "Poster Concept", "Book Cover Design",  

            # Illustration & Graphic Elements  
            "Illustration Concept", "Icon Set", "Custom Vector Art", "Infographic Layout",  
            "Social Media Graphics", "Sticker Pack", "3D Design Concept",  

            # Advanced & Systematic Design  
            "Design System", "UI Component Library", "Grid System Guide", "Motion Design Concept",  
            "Microinteraction Design", "Design Tokens"  
        ],

        
        "Development": [
            # Software Planning & Architecture  
            "Development Plan", "Code Architecture", "Tech Stack", "System Architecture",  
            "Software Requirements", "Scalability Strategy", "Performance Optimization Guide",  

            # Coding & Documentation  
            "Code Review", "Code Snippet", "Code Refactoring", "Algorithm", "API Documentation",  
            "Database Schema", "Technical Spec", "Feature Spec",  

            # Testing & Debugging  
            "Testing Strategy", "Unit Test Case", "Bug Report", "Debugging Guide",  
            "Security Best Practices",  

            # DevOps & Deployment  
            "CI/CD Pipeline Plan", "Infrastructure as Code", "Deployment Strategy",  
            "Server Configuration", "Monitoring & Logging Setup",  

            # Web & Mobile Development  
            "Frontend Framework Guide", "Backend API Design", "Full-Stack Development Plan",  
            "Responsive Web Design Principles", "Cross-Platform Development Strategy",  

            # AI & Data Science Development  
            "ML Model Deployment Guide", "Data Pipeline Design", "AI Ethics & Bias Mitigation",  
            "Deep Learning Framework Selection", "Edge AI Development Plan"  
        ],

        "Marketing": [
            # Strategy & Planning  
            "Marketing Plan", "Campaign Brief", "Growth Hack", "Product Launch Strategy", "Rebranding Strategy",  
            "Customer Journey", "Competitive Positioning", "Go-to-Market Plan", "Market Segmentation",  

            # Advertising & Promotion  
            "Ad Copy", "Landing Page", "Promotion", "Social Media Ad Script", "Native Advertising Copy",  
            "Influencer Marketing Strategy", "Referral Marketing Plan", "Affiliate Marketing Plan",  

            # Content & Engagement  
            "Email Campaign", "Newsletter Strategy", "Content Calendar", "SEO-Optimized Blog Strategy",  
            "Video Marketing Plan", "Podcast Promotion Plan", "Webinar Outline",  

            # Sales & Conversion Optimization  
            "Conversion Strategy", "Sales Script", "Lead Magnet Idea", "Customer Retention Plan",  
            "Call-to-Action Optimization", "Upselling & Cross-Selling Strategy",  

            # Branding & Messaging  
            "Value Proposition", "USP", "Elevator Pitch", "Messaging Framework", "Brand Storytelling",  
            "Emotional Branding Guide", "Tagline & Slogan Generator",  

            # Market Research & Analytics  
            "Customer Persona", "A/B Testing Plan", "Performance Metrics Dashboard", "Marketing ROI Calculator",  
            "Competitor Analysis", "Consumer Trend Report", "Pricing Psychology Strategy"  
        ],
        "Finance": [
            # Budgeting & Financial Planning  
            "Budget Plan", "Expense Report", "Savings Plan", "Personal Finance Tracker", "Cash Flow Management",  
            "Debt Management", "Cost Reduction Strategy", "Emergency Fund Planning",  

            # Investment & Business Finance  
            "Investment Strategy", "Stock Market Analysis", "Wealth Management Plan", "Equity Distribution",  
            "Fundraising Strategy", "Investor Pitch Financials", "Mergers & Acquisitions Strategy",  

            # Revenue & Profit Optimization  
            "Revenue Forecast", "Profit Optimization", "Monetization Strategy", "Pricing Model",  
            "Financial Risk Assessment", "Financial Growth Projection",  

            # Tax & Compliance  
            "Tax Strategy", "Legal & Compliance Checklist", "Audit Preparation Guide",  

            # Retirement & Long-Term Planning  
            "Retirement Plan", "Passive Income Strategy", "Estate Planning Guide",  

            # Financial Education & Analysis  
            "Financial Model", "Financial Education", "Personal Net Worth Calculation",  
            "Business Valuation Analysis", "Break-Even Analysis"  
        ],

        "Health": [
            # General Wellness & Lifestyle  
            "Wellness Plan", "Self-Care Routine", "Health Goal", "Stress Management", "Work-Life Balance Guide",  
            "Sleep Improvement", "Mindfulness Exercise", "Daily Energy Optimization",  

            # Diet & Nutrition  
            "Diet Plan", "Meal Prep Guide", "Nutrition Guide", "Hydration Plan", "Superfood Recommendations",  
            "Vitamins & Supplements Guide", "Intermittent Fasting Plan",  

            # Fitness & Exercise  
            "Fitness Routine", "Home Workout Plan", "Gym Training Plan", "Yoga Sequence",  
            "Cardio vs Strength Training Guide", "Running Plan", "HIIT Routine",  

            # Mental Health & Recovery  
            "Mental Health Check-In", "Meditation Script", "Anxiety & Stress Relief Guide",  
            "Cognitive Behavioral Therapy (CBT) Techniques", "Journaling Prompts for Mental Clarity",  

            # Symptom & Medical Insights  
            "Symptom Analysis", "Health Tracker", "Chronic Illness Management Plan",  
            "Medical Information", "Preventive Healthcare Checklist",  

            # Specialized & Alternative Health  
            "Holistic Healing Plan", "Ayurveda & Herbal Remedies", "Acupuncture & Alternative Medicine",  
            "Post-Injury Recovery Plan", "Rehabilitation Strategy"  
        ],
        "Legal": [
            # Legal Drafting & Contracts  
            "Contract Template", "Legal Letter", "Legal Response", "Regulatory Filing",  
            "Terms of Service", "Privacy Statement", "Disclaimer", "Copyright Notice",  

            # Business & Intellectual Property (IP)  
            "IP Strategy", "Trademark Filing Guide", "Patent Summary", "Intellectual Property Protection Plan",  
            "Business Compliance Checklist", "Non-Disclosure Agreement (NDA)",  

            # Legal Analysis & Dispute Resolution  
            "Legal Analysis", "Legal Research", "Dispute Resolution", "Regulatory Compliance Audit",  
            "Risk Assessment & Mitigation", "Case Law Review",  

            # Consumer & Personal Legal Support  
            "Employment Contract Review", "Tenant Rights & Rental Agreement",  
            "Small Claims Court Guide", "Legal Defense Outline", "Cyber Law & Data Protection Strategy"  
        ],

        "Event": [
            # Event Planning & Logistics  
            "Event Plan", "Event Budget", "Venue Description", "Event Schedule", "Guest List",  
            "Conference Agenda", "Corporate Event Strategy", "Virtual Event Setup",  

            # Invitations & Announcements  
            "Invitation", "Save-the-Date Announcement", "Event Reminder Message", "RSVP Confirmation",  

            # Speeches & Messages  
            "Wedding Speech", "Toast", "Thank You Note", "Opening Ceremony Speech",  
            "Farewell Speech", "Award Ceremony Speech", "Eulogy",  

            # Event Content & Marketing  
            "Event Marketing", "Social Media Event Promotion", "Press Release for Event",  
            "Sponsorship Proposal", "Hashtag Strategy for Events",  

            # Entertainment & Activities  
            "Party Theme", "Entertainment Plan", "Catering Menu", "Music Playlist Suggestion",  
            "Photo Booth Ideas", "Games & Icebreakers", "Team Building Activity",  

            # Specialized Event Plans  
            "Wedding Planning Guide", "Birthday Party Plan", "Baby Shower Ideas",  
            "Graduation Ceremony Plan", "Charity Fundraiser Outline", "Festival & Fair Planning"  
        ],
        
        "Relationships": [
            # Romantic Relationships  
            "Love Letter", "Apology Letter", "Breakup Letter", "Anniversary Note", "Proposal Message",  
            "Long-Distance Relationship Advice", "Heartfelt Compliment Generator",  

            # Friendships & Social Connections  
            "Friendship Message", "Birthday Message", "Congratulations Note", "Supportive Message",  
            "Reconnection Message", "Group Chat Icebreaker",  

            # Family & Personal Relationships  
            "Family Communication Guide", "Parenting Advice", "Sibling Bonding Ideas",  
            "Holiday Greeting Message", "Personalized Family Storytelling",  

            # Professional & Networking  
            "Networking Message", "LinkedIn Connection Request", "Mentor Outreach Email",  
            "Professional Thank You Note", "Colleague Appreciation Message",  

            # Emotional Support & Conflict Resolution  
            "Relationship Advice", "Conflict Resolution Guide", "Condolence Note",  
            "Forgiveness Letter", "Active Listening Tips", "Empathy Coaching"  
        ],
 
        "Industry": [
            # Industry Research & Market Analysis  
            "Industry Analysis", "Sector Trend", "Market Forecast", "Industry Report",  
            "Competitive Landscape", "Market Share Analysis", "Economic Impact Report",  

            # Regulations & Compliance  
            "Regulatory Impact", "Government Policy Update", "Industry Compliance Checklist",  
            "International Trade Law Overview", "ESG (Environmental, Social, and Governance) Compliance",  

            # Technology & Innovation  
            "Technology Adoption", "AI & Automation in Industry", "Sustainability Trends",  
            "Industry 4.0 Strategy", "Blockchain Use Cases in Industry", "Green Energy Impact",  

            # Business Growth & Strategy  
            "Industry Disruption", "Vertical Strategy", "Expansion & Globalization Plan",  
            "Mergers & Acquisitions Insight", "Franchise Model Analysis",  

            # Supply Chain & Distribution  
            "Supply Chain Optimization", "Logistics Strategy", "Reshoring vs Offshoring Analysis",  
            "Distribution Channel Strategy", "Inventory Management Trends",  

            # Partnerships & Networking  
            "Industry Partnership", "Joint Venture Feasibility", "Trade Association Guide",  
            "Industry Event Planning", "Corporate Sponsorship Strategy",  

            # Future Trends & Predictions  
            "Emerging Market Analysis", "Future of Work Report", "Disruptive Startup Watchlist",  
            "AI Impact on Industry", "Space Economy Trends", "Cybersecurity Risk Forecast"  
        ]
 }
    
    # Combine base and extended categories
    all_categories = {**base_categories, **extended_categories}
    
    # Multipliers to expand each tool category (adjectival prefixes)
# Practical Tool Multipliers

    # Descriptive Qualifiers: Focus on Actual Capabilities
    tool_multipliers = [
        "AI-Powered", "Data-Driven", "Adaptive", "Intelligent", 
        "Automated", "Precision", "Workflow", "Strategic", 
        "Analytical", "Efficient", "Scalable", "Integrated"
    ]

    # Tool Format Descriptors: Emphasize Functional Purpose
    format_multipliers = [
        "Engine", "Platform", "Framework", "Toolkit", 
        "Accelerator", "Assistant", "Optimizer", "Analyzer", 
        "Manager", "Generator", "Orchestrator", "Solution"
    ]
    
    # Generate expanded tools by combining base tools with multipliers
    expanded_categories = {}
    for category, base_tools in all_categories.items():
        expanded_tools = []
        for base_tool in base_tools:
            # Add the original tool with common suffixes
            for format_mult in format_multipliers[:3]:  # Limit to top 3 formats
                expanded_tools.append(f"{base_tool} {format_mult}")
            
            # Add tools with multipliers (limit to reduce overwhelming options)
            for mult in tool_multipliers[:5]:  # Limit to top 5 multipliers
                expanded_tools.append(f"{mult} {base_tool}")
        
        expanded_categories[category] = expanded_tools
    
    # Create specialized industry categories (new areas)
    specialized_industries = {
        "Healthcare": ["Patient Care", "Medical Record", "Clinical Trial", "Health Assessment", "Treatment Plan",
                      "Healthcare Policy", "Medical Research", "Patient Education", "Telehealth", "Health Insurance",
                      "Medical Device", "Healthcare Compliance", "Medical Diagnosis", "Patient Experience", "Wellness Program"],
        "E-commerce": [
            # Product & Inventory Management  
            "Product Listing Optimization", "Dynamic Pricing Strategy", "Inventory Forecasting",  
            "Dropshipping Business Plan", "Product Lifecycle Management",  

            # Customer Experience & Retention  
            "Customer Review Strategy", "Shopping Experience Enhancement", "Customer Support AI Chatbot",  
            "Loyalty Program Development", "Subscription Model Strategy",  

            # Marketing & Sales Growth  
            "E-commerce Copywriting", "SEO for Product Pages", "Email Marketing for E-commerce",  
            "Influencer Marketing for Online Stores", "Flash Sale Campaign Strategy",  

            # Logistics & Operations  
            "Shipping Policy Optimization", "Return & Refund Policy", "Cross-Border E-commerce Strategy",  
            "Last-Mile Delivery Innovation", "Omnichannel Retailing Strategy",  

            # E-commerce Business Development  
            "Marketplace Strategy", "D2C (Direct-to-Consumer) Business Plan", "B2B E-commerce Expansion",  
            "Mobile Commerce (M-commerce) Strategy", "Live Shopping & Video Commerce"  
        ],
        
   
        "Real Estate": ["Property Description", "Market Analysis", "Investment Property", "Rental Analysis", "Home Staging",
                       "Property Marketing", "Neighborhood Guide", "Real Estate Listing", "Agent Bio", "Mortgage Information",
                       "Home Inspection", "Lease Agreement", "Property Management", "HOA Communication", "Commercial Lease"],
        
        "Sustainability": [
            # Environmental Impact & Reporting  
            "Environmental Impact Assessment", "Sustainability Report", "Carbon Footprint Reduction Plan",  
            "Greenhouse Gas Emissions Analysis", "Corporate Sustainability Strategy",  

            # Sustainable Business & Economy  
            "ESG (Environmental, Social, Governance) Strategy", "Circular Economy Implementation",  
            "Sustainable Supply Chain Practices", "Green Business Certification Guide",  

            # Renewable Energy & Conservation  
            "Energy Efficiency Strategy", "Solar Energy Adoption Plan", "Water Conservation Policy",  
            "Waste Reduction & Recycling Strategy", "Carbon Offsetting Guide",  

            # Climate Change & Social Responsibility  
            "Climate Action Plan", "Biodiversity Conservation Plan", "Environmental Justice Advocacy",  
            "Sustainable Urban Development", "Green Policy Proposal",  

            # Sustainable Design & Construction  
            "Sustainable Architecture Guide", "Eco-Friendly Building Materials", "LEED Certification Process",  
            "Net-Zero Energy Homes", "Smart City & Urban Sustainability",  

            # Consumer & Lifestyle Sustainability  
            "Low-Waste Living Guide", "Ethical Consumerism Strategy", "Sustainable Fashion Practices",  
            "Plant-Based Diet & Environmental Impact", "Green Travel & Ecotourism"  
        ],

        "Technology": [
            # **Software Development & Engineering**  
            "Tech Specification", "Software Release Plan", "Product Roadmap", "API Documentation",  
            "Code Optimization Strategy", "Software Architecture Design", "DevOps Best Practices",  

            # **User Experience & Support**  
            "User Guide & Documentation", "Tech Support Response", "Troubleshooting Workflow",  
            "Accessibility & Inclusive Tech Design", "Onboarding & Training Guide",  

            # **Cybersecurity & Data Protection**  
            "Cybersecurity Framework", "Data Privacy Policy", "Threat Detection Strategy",  
            "Ransomware Protection Plan", "Network Security Guide",  

            # **Cloud & Infrastructure**  
            "Cloud Migration Strategy", "Multi-Cloud vs Hybrid Cloud Analysis", "Serverless Architecture Guide",  
            "IT Infrastructure Roadmap", "Disaster Recovery & Backup Plan",  

            # **AI, Machine Learning & Emerging Tech**  
            "AI Integration Strategy", "Machine Learning Model Deployment", "AI Ethics & Bias Mitigation",  
            "Quantum Computing Roadmap", "Blockchain Use Cases",  

            # **Hardware & IoT**  
            "Hardware Design Guide", "Embedded Systems Optimization", "IoT Security Strategy",  
            "Edge Computing Implementation", "Smart Home & Automation Guide",  

            # **Enterprise IT & Digital Transformation**  
            "IT Strategy & Governance", "Digital Transformation Roadmap", "Tech Integration Best Practices",  
            "Enterprise Software Selection", "Legacy System Modernization",  

            # **Tech Evaluation & Market Trends**  
            "Tech Evaluation & Feasibility Study", "Emerging Tech Trend Analysis", "Gartner Hype Cycle Insights",  
            "Future of Work & Automation", "5G & Connectivity Strategy"  
        ]
  }
    
    # Add specialized industries to the expanded categories
    for industry, tools in specialized_industries.items():
        industry_tools = []
        for tool in tools:
            # Add the original tool with common suffixes
            for format_mult in format_multipliers[:3]:  # Limit to top 3 formats
                industry_tools.append(f"{tool} {format_mult}")
            
            # Add tools with multipliers (limit to reduce overwhelming options)
            for mult in tool_multipliers[:3]:  # Limit to top 3 multipliers
                industry_tools.append(f"{mult} {tool}")
        
        expanded_categories[industry] = industry_tools
    
    # Create flat list of all tools
    all_tools = []
    for category, tools in expanded_categories.items():
        all_tools.extend(tools)
    
    return all_tools, expanded_categories


# Add this to your imports
import re
from collections import Counter

# Add this function after your other functions
def analyze_content(text):
    # Basic text analysis
    word_count = len(re.findall(r'\w+', text))
    sentence_count = len(re.findall(r'[.!?]+', text)) + 1
    paragraph_count = len(text.split('\n\n'))
    
    # Readability metrics
    words = re.findall(r'\w+', text.lower())
    word_lengths = [len(word) for word in words]
    avg_word_length = sum(word_lengths) / len(word_lengths) if word_lengths else 0
    
    # Sentiment indicators (simplified)
    positive_words = ['good', 'great', 'excellent', 'best', 'positive', 'happy', 'wonderful', 'amazing']
    negative_words = ['bad', 'worst', 'terrible', 'negative', 'poor', 'awful', 'horrible']
    
    positive_count = sum(1 for word in words if word in positive_words)
    negative_count = sum(1 for word in words if word in negative_words)
    
    sentiment = "Positive" if positive_count > negative_count else "Negative" if negative_count > positive_count else "Neutral"
    
    # Most common words
    common_words = Counter(words).most_common(5)
    
    return {
        "word_count": word_count,
        "sentence_count": sentence_count,
        "paragraph_count": paragraph_count,
        "avg_word_length": round(avg_word_length, 2),
        "sentiment": sentiment,
        "common_words": common_words
    }
    
# Get all tools and categories
ai_tools, tool_categories = generate_ai_tools()

# Function to load prompt templates (dynamic generation)
@st.cache_data
def load_prompt_templates():
    templates = {}
    for tool in ai_tools:
        templates[tool] = f"Generate content using the '{tool}' feature for: {{prompt}}. Ensure the output is high quality, relevant, and tailored to the user's needs."
    return templates

# Function to generate content with AI
def generate_ai_content(prompt, api_key, model_name):
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        with st.spinner("🔮 AI is working its magic..."):
            generation_config = {"temperature": 0.7, "top_p": 0.95, "top_k": 40, "max_output_tokens": 8192}
            response = model.generate_content(prompt, generation_config=generation_config)
            return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# Function to save content to history
def save_to_history(tool_name, prompt, output):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.session_state.history.insert(0, {"timestamp": timestamp, "tool": tool_name, "prompt": prompt, "output": output})
    if len(st.session_state.history) > 20:
        st.session_state.history = st.session_state.history[:20]

# Sidebar for API configuration
with st.sidebar:
    st.image("aichip.jpg")
    st.markdown("### 🔑 API Configuration")
    api_key = st.text_input("Enter Google Gemini API Key:", type="password", value=st.session_state.api_key, help="Your API key is stored in the session and not saved.")
    if api_key: st.session_state.api_key = api_key
    
    # Model selection with Gemini models
    st.session_state.api_model = st.selectbox(
        "Select AI Model:",
        ["gemini-2.0-flash", "gemini-1.5-flash", "gemini-1.5-pro","gemini-2.0-flash-lite","gemini-2.0-pro-exp-02-05",
"gemini-2.0-flash-thinking-exp-01-21","gemini-2.5-pro-exp-03-25","gemini-1.5-flash-8b"],
        index=0
    )
    
    st.markdown("---")
    search_term = st.text_input("🔍 Search for tools:")
    filtered_tools = [tool for tool in ai_tools if search_term.lower() in tool.lower()] if search_term else ai_tools
    
    # Statistics
    # Statistics
    st.sidebar.markdown("### 📊 Stats")
    st.sidebar.markdown(f"**Total Tools:** {len(ai_tools)}")
    st.sidebar.markdown(f"**Categories:** {len(tool_categories)}")
    
    # History button in sidebar
    st.sidebar.markdown("---")
    show_history = st.sidebar.button("📜 View History")
    clear_history = st.sidebar.button("🗑️ Clear History")
    if clear_history:
        st.session_state.history = []
        st.sidebar.success("History cleared!")

# Main content area
st.title("🧠 Ultimate AI Creator Hub")
st.markdown("Create amazing content with AI-powered tools - all in one place!")

# Welcome message
if not st.session_state.api_key:
    st.info("👋 Welcome! Please enter your API key in the sidebar to get started.")

# Display history if requested
if show_history and 'history' in st.session_state and len(st.session_state.history) > 0:
    st.header("📜 Content History")
    for i, item in enumerate(st.session_state.history):
        with st.expander(f"{item['timestamp']} - {item['tool']}"):
            st.markdown("**Prompt:**")
            st.markdown(f"```\n{item['prompt']}\n```")
            st.markdown("**Output:**")
            st.markdown(item['output'])
    st.markdown("---")

# Load prompt templates if needed
if not st.session_state.prompt_templates or len(st.session_state.prompt_templates) == 0:
    st.session_state.prompt_templates = load_prompt_templates()

# Tool Selection Section
st.header("🛠️ Select Your Creation Tool")

# Apply Scrollable Tabs CSS
st.markdown(
    """
    <style>
    div[data-testid="stTabs"] {
        overflow-x: auto !important;
        white-space: nowrap !important;
        display: flex;
        flex-wrap: nowrap;
    }
    div[data-testid="stTabs"] button {
        flex: 1 1 auto;
        min-width: 150px;  /* Adjust tab width */
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.markdown('<div class="scroll-container">', unsafe_allow_html=True)
# Close the Scrollable Container

tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10, tab11, tab12, tab13, tab14, tab15 = st.tabs([
    "📋 Categories", "🔍 Search", 
    "📚 Research", "🤖 Chat", 
    "🌍 Translate", "⚡ Code", 
    "📊 Insights", "🎤 Interview",
    "📧 Email Assistant", "📊 Spreadsheet",
    "🎬 Podcast", "🎯 Learning Path",
    "📝 Meeting Minutes", "🚀 Startup Validator",
    "🔮 AI Search"
])





# Wrap Everything Below Tabs in a Scrollable Container

st.markdown("</div>", unsafe_allow_html=True)

with tab1:
    selected_category = st.selectbox("Choose a category:", list(tool_categories.keys()))
    
    # Only show tools from selected category
    if selected_category:
        tools_in_category = tool_categories[selected_category]
        st.markdown(f"### {selected_category} Tools ({len(tools_in_category)})")
        
        # Create grid layout for tools
        cols = st.columns(3)
        for i, tool in enumerate(tools_in_category):
            with cols[i % 3]:
                if st.button(tool, key=f"cat_{tool}"):
                    st.session_state.selected_tool = tool

with tab2:
    if search_term:
        st.markdown(f"### Search Results for '{search_term}' ({len(filtered_tools)})")
        
        # Create grid layout for search results
        cols = st.columns(3)
        for i, tool in enumerate(filtered_tools[:30]):  # Limit to 30 results
            with cols[i % 3]:
                if st.button(tool, key=f"search_{tool}"):
                    st.session_state.selected_tool = tool
    else:
        st.info("Enter a search term in the sidebar to find specific tools.")


with tab3:
    st.header("📚 AI Research Assistant")

    # Paper Summarization Section - ENHANCED
    st.subheader("📑 Research Paper Summarization")
    uploaded_file = st.file_uploader("Upload Research Paper (PDF)", type="pdf")
    
    if uploaded_file:
        pdf_text = extract_text_from_pdf(uploaded_file)
        
        # Add advanced PDF extraction options
        extraction_options = st.expander("Advanced Extraction Options")
        with extraction_options:
            extract_figures = st.checkbox("Extract Figures", value=False)
            extract_tables = st.checkbox("Extract Tables", value=False)
            extract_references = st.checkbox("Extract References", value=True)
            max_text_length = st.slider("Maximum Text Length (chars)", 1000, 20000, 6000)
        
        st.text_area("Extracted Text:", pdf_text[:max_text_length], height=500)
        
        summary_options = st.expander("Summary Options")
        with summary_options:
            summary_type = st.radio("Summary Type", 
                ["Concise (1 paragraph)", "Standard (3-5 paragraphs)", "Detailed (comprehensive)"])
            focus_areas = st.multiselect("Focus Areas", 
                ["Methodology", "Results", "Conclusions", "Background", "Limitations", "Future Work"])
            academic_level = st.select_slider("Academic Level", 
                ["Undergraduate", "Graduate", "Expert"])
        
        if st.button("Summarize Paper ✨"):
            # Enhanced prompt with options
            focus_str = ", ".join(focus_areas) if focus_areas else "all sections"
            summary_prompt = f"""Summarize this research paper {summary_type}. 
            Focus on {focus_str} at an {academic_level} level:
            
            {pdf_text[:max_text_length]}"""
            
            with st.spinner("Generating summary..."):
                summary = generate_ai_content(summary_prompt, st.session_state.api_key, st.session_state.api_model)
            
            st.success("📑 AI Summary:")
            st.write(summary)

            # Add export options
            col1, col2 = st.columns(2)
            with col1:
                st.download_button("Download Summary", summary, "paper_summary.txt")
            with col2:
                st.button("Copy to Clipboard", on_click=lambda: st.write("<script>navigator.clipboard.writeText(`" + summary.replace("`", "\\`") + "`);</script>", unsafe_allow_html=True))

    # Citation Generator - ENHANCED
    st.subheader("📖 Citation Generator")
    
    citation_tabs = st.tabs(["Text Input", "PDF Upload", "DOI Lookup"])
    
    with citation_tabs[0]:
        citation_input = st.text_area("Paste Reference Text:")
        
    with citation_tabs[1]:
        citation_pdf = st.file_uploader("Upload paper for citation extraction", type="pdf")
        if citation_pdf:
            citation_input = extract_text_from_pdf(citation_pdf)
            st.success("PDF processed for citations")
            
    with citation_tabs[2]:
        doi_input = st.text_input("Enter DOI (e.g., 10.1038/nature12373):")
        if doi_input and st.button("Fetch DOI Metadata"):
            # Implement DOI lookup API call
            st.session_state.citation_input = f"DOI: {doi_input}"
            citation_input = st.session_state.citation_input
            st.success(f"Retrieved metadata for DOI: {doi_input}")
    
    citation_format = st.selectbox("Citation Format:", 
        ["IEEE", "APA", "MLA", "Chicago", "Harvard", "BibTeX", "RIS"])
    
    if st.button("Generate Citations"):
        citations = extract_references(citation_input)
        citation_prompt = f"Convert these references into {citation_format} format:\n{citations}"
        
        with st.spinner("Formatting citations..."):
            formatted_citations = generate_ai_content(citation_prompt, st.session_state.api_key, st.session_state.api_model)
        
        st.success("📜 Formatted Citations:")
        st.code(formatted_citations)
        
        # Add copy and export options
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download Citations", formatted_citations, f"citations_{citation_format.lower()}.txt")
        with col2:
            if citation_format == "BibTeX":
                st.download_button("Download BibTeX", formatted_citations, "references.bib")

    # Research Idea Expansion - ENHANCED
    st.subheader("🔬 Research Proposal Generator")
    
    research_topic = st.text_input("Enter Research Idea:")
    
    proposal_options = st.expander("Proposal Options")
    with proposal_options:
        proposal_type = st.selectbox("Proposal Type", 
            ["Brief Concept", "Conference Abstract", "Full Research Proposal", "Grant Application"])
        
        discipline = st.selectbox("Academic Discipline", 
            ["Computer Science", "Biology", "Chemistry", "Physics", "Psychology", 
             "Sociology", "Economics", "Engineering", "Medicine", "Environmental Science", "Other"])
        
        if discipline == "Other":
            custom_discipline = st.text_input("Specify Discipline:")
            if custom_discipline:
                discipline = custom_discipline
        
        methodologies = st.multiselect("Preferred Methodologies", 
            ["Quantitative", "Qualitative", "Mixed Methods", "Experimental", 
             "Observational", "Literature Review", "Meta-Analysis", "Simulation"])
        
        novelty_slider = st.slider("Novelty Level", 1, 10, 7, 
            help="1 = Incremental research, 10 = Groundbreaking approach")
        
        word_limit = st.number_input("Word Count Target", 300, 5000, 1000)
    
    if st.button("Generate Proposal 🚀"):
        # Build advanced prompt with all parameters
        methods_str = ", ".join(methodologies) if methodologies else "any appropriate"
        
        expansion_prompt = f"""Develop a {proposal_type} for: {research_topic}
        
        Academic field: {discipline}
        Methodologies to consider: {methods_str}
        Novelty level (1-10): {novelty_slider}
        Target length: ~{word_limit} words
        
        Include: research question, background, methodology, expected outcomes, and significance.
        """
        
        with st.spinner("Creating research proposal..."):
            research_proposal = generate_ai_content(expansion_prompt, st.session_state.api_key, st.session_state.api_model)
        
        st.success("📄 AI-Generated Research Proposal:")
        st.write(research_proposal)
        
        # Export options
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("Download Proposal", research_proposal, "research_proposal.docx")
        with col2:
            st.button("Save to Projects", help="Save this proposal to your projects collection")


with tab4:
    st.header("🤖 AI Chatbot with Universal File Upload")
    
    # Import required libraries
    import pandas as pd
    import matplotlib.pyplot as plt
    import plotly.express as px
    from collections import Counter
    import re
    from io import StringIO
    
    # Try to import wordcloud or handle missing dependency
    try:
        from wordcloud import WordCloud, STOPWORDS
    except ImportError:
        st.warning("WordCloud is not installed. Install with: pip install wordcloud")
        # Create dummy functions/classes to prevent errors
        class WordCloud:
            def __init__(self, **kwargs):
                pass
            def generate(self, text):
                pass
        STOPWORDS = set()
    
    # Try to import nltk or handle missing dependency
    try:
        import nltk
        from nltk.util import ngrams
    except ImportError:
        st.warning("NLTK is not installed. Install with: pip install nltk")
        # Create dummy functions to prevent errors
        def ngrams(text, n):
            return []

    # File uploader
    uploaded_file = st.file_uploader("Upload a file (PDF, DOCX, CSV, TXT, Audio)", 
                                    type=["pdf", "docx", "csv", "txt", "mp3", "wav"])

    extracted_text = ""  # Store extracted text
    df = None  # Store dataframe for visualizations

    # Process uploaded file
    if uploaded_file:
        file_type = uploaded_file.type

        if file_type == "application/pdf":
            extracted_text = extract_text_from_pdf(uploaded_file)
        elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            extracted_text = extract_text_from_docx(uploaded_file)
        elif file_type in ["text/plain"]:
            extracted_text = extract_text_from_txt(uploaded_file)
            try:
                # Try to parse as CSV in case it's a CSV saved as TXT
                if extracted_text.strip():  # Check if text is not empty
                    df = pd.read_csv(StringIO(extracted_text), sep=None, engine='python')
                else:
                    words = re.findall(r'\w+', extracted_text.lower())
                    word_freq = Counter(words)
            except Exception as e:
                st.warning(f"Could not parse as CSV: {str(e)}")
                # If not parseable as CSV, prepare for text visualizations
                words = re.findall(r'\w+', extracted_text.lower())
                word_freq = Counter(words)
        elif file_type in ["text/csv"]:
            extracted_text = extract_text_from_csv(uploaded_file)
            try:
                # Reset file pointer to beginning before reading
                uploaded_file.seek(0)
                df = pd.read_csv(uploaded_file)
                if df.empty:
                    st.warning("The CSV file appears to be empty")
            except Exception as e:
                st.warning(f"Error parsing CSV: {str(e)}")
                # Handle as text if CSV parsing fails
                words = re.findall(r'\w+', extracted_text.lower())
                word_freq = Counter(words)
        elif file_type in ["image/png", "image/jpeg"]:
            extracted_text = extract_text_from_image(uploaded_file)
        elif file_type in ["audio/mpeg", "audio/wav"]:
            extracted_text = extract_text_from_audio(uploaded_file)

        st.success("✅ File processed successfully!")
        st.text_area("Extracted Content:", extracted_text[:2000], height=200)  # Preview first 2000 characters

        # Visualization section for TXT and CSV
        if file_type in ["text/plain", "text/csv"] and (df is not None or 'word_freq' in locals()):
            st.subheader("📊 Data Visualizations")
            
            # Create tabs for different visualization categories
            viz_tabs = st.tabs(["Basic Stats", "Distributions", "Correlations", "Time Series", "Text Analysis"])
            
            # CSV-specific visualizations
            if df is not None:
                with viz_tabs[0]:  # Basic Stats
                    st.write("### Data Overview")
                    st.dataframe(df.describe())
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.write("### Missing Values")
                        st.bar_chart(df.isnull().sum())
                    with col2:
                        st.write("### Data Types")
                        st.write(pd.DataFrame(df.dtypes, columns=['Data Type']))
                
                with viz_tabs[1]:  # Distributions
                    st.write("### Distributions")
                    
                    # Determine numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
                    if numeric_cols:
                        selected_col = st.selectbox("Select column for histogram:", numeric_cols)
                        fig = px.histogram(df, x=selected_col, marginal="box")
                        st.plotly_chart(fig, use_container_width=True)
                        
                        # Distribution comparison
                        if len(numeric_cols) >= 2:
                            cols_to_compare = st.multiselect("Select columns to compare distributions:", 
                                                           numeric_cols, default=numeric_cols[:2])
                            if cols_to_compare:
                                fig = px.box(df, y=cols_to_compare)
                                st.plotly_chart(fig, use_container_width=True)
                    
                    # Categorical distributions
                    cat_cols = df.select_dtypes(include=['object']).columns.tolist()
                    if cat_cols:
                        selected_cat = st.selectbox("Select categorical column:", cat_cols)
                        top_n = st.slider("Top N categories:", 5, 20, 10)
                        value_counts = df[selected_cat].value_counts().head(top_n)
                        fig = px.bar(x=value_counts.index, y=value_counts.values)
                        st.plotly_chart(fig, use_container_width=True)
                
                with viz_tabs[2]:  # Correlations
                    st.write("### Correlations")
                    if len(numeric_cols) >= 2:
                        try:
                            corr_method = st.radio("Correlation method:", ["Pearson", "Spearman"])
                            corr = df[numeric_cols].corr(method=corr_method.lower())
                            
                            fig = px.imshow(corr, text_auto=True, color_continuous_scale='RdBu_r')
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Scatter plot for two variables
                            st.write("### Scatter Plot")
                            col1, col2 = st.columns(2)
                            with col1:
                                x_col = st.selectbox("X-axis:", numeric_cols, index=0)
                            with col2:
                                y_col = st.selectbox("Y-axis:", numeric_cols, index=min(1, len(numeric_cols)-1))
                            
                            color_col = st.selectbox("Color by (optional):", ["None"] + df.columns.tolist())
                            if color_col == "None":
                                fig = px.scatter(df, x=x_col, y=y_col)
                            else:
                                fig = px.scatter(df, x=x_col, y=y_col, color=color_col)
                            st.plotly_chart(fig, use_container_width=True)
                        except Exception as e:
                            st.error(f"Error creating correlation visualizations: {str(e)}")
                    else:
                        st.info("Need at least 2 numeric columns for correlation analysis.")
                
                with viz_tabs[3]:  # Time Series
                    st.write("### Time Series Analysis")
                    
                    # Check for date columns
                    date_cols = []
                    for col in df.columns:
                        try:
                            pd.to_datetime(df[col])
                            date_cols.append(col)
                        except:
                            pass
                    
                    if date_cols:
                        date_col = st.selectbox("Select date column:", date_cols)
                        
                        # Convert to datetime
                        df_ts = df.copy()
                        df_ts[date_col] = pd.to_datetime(df_ts[date_col])
                        
                        # Select value column
                        value_col = st.selectbox("Select value column:", numeric_cols)
                        
                        # Resample options
                        resample_options = ['Original', 'Daily', 'Weekly', 'Monthly', 'Quarterly', 'Yearly']
                        resample = st.selectbox("Resample frequency:", resample_options)
                        
                        # Prepare data based on resampling
                        if resample == 'Original':
                            plot_data = df_ts.sort_values(by=date_col)
                        else:
                            # Map selection to pandas frequency
                            freq_map = {'Daily': 'D', 'Weekly': 'W', 'Monthly': 'M', 
                                       'Quarterly': 'Q', 'Yearly': 'Y'}
                            df_ts = df_ts.set_index(date_col)
                            plot_data = df_ts[value_col].resample(freq_map[resample]).mean().reset_index()
                        
                        # Line chart
                        fig = px.line(
                            plot_data, 
                            x=date_col if resample == 'Original' else 'index', 
                            y=value_col
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.info("No date columns detected. To perform time series analysis, ensure one of your columns contains dates.")
                
                with viz_tabs[4]:  # Text Analysis
                    st.write("### Text Analysis")
                    
                    # Check for text columns
                    text_cols = df.select_dtypes(include=['object']).columns.tolist()
                    
                    if text_cols:
                        text_col = st.selectbox("Select text column:", text_cols)
                        
                        # Word cloud
                        st.write("#### Word Cloud")
                        
                        # Join all text
                        all_text = " ".join(df[text_col].dropna().astype(str))
                        
                        if all_text.strip():
                            # Check if there are words after filtering
                            words = re.findall(r'\w+', all_text.lower())
                            stop_words = set(STOPWORDS)
                            filtered_words = [w for w in words if w not in stop_words and len(w) > 2]
                            
                            if filtered_words:  # Add this check
                                # Generate word cloud
                                wc = WordCloud(width=800, height=400, 
                                              background_color='white', 
                                              colormap='viridis', 
                                              max_words=200)
                                wc.generate(all_text)
                                
                                # Display
                                fig, ax = plt.subplots(figsize=(10, 5))
                                ax.imshow(wc, interpolation='bilinear')
                                ax.axis('off')
                                st.pyplot(fig)
                            else:
                                st.info("Not enough meaningful words to generate a word cloud after filtering.")
                            
                            # Word frequency
                            st.write("#### Top Words")
                            stop_words = set(STOPWORDS)
                            words = re.findall(r'\w+', all_text.lower())
                            word_freq = Counter([w for w in words if w not in stop_words and len(w) > 2])
                            
                            if word_freq:
                                top_words = pd.DataFrame(word_freq.most_common(20), columns=['Word', 'Count'])
                                fig = px.bar(top_words, x='Word', y='Count')
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info("No significant words found for frequency analysis.")
                    else:
                        st.info("No text columns detected for text analysis.")
            
            # Text-specific visualizations for TXT files
            elif 'word_freq' in locals():
                with viz_tabs[0]:  # Basic Stats
                    st.write("### Text Statistics")
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Words", len(words))
                    with col2:
                        st.metric("Unique Words", len(word_freq))
                    with col3:
                        st.metric("Total Characters", len(extracted_text))
                
                with viz_tabs[4]:  # Text Analysis
                    st.write("### Text Analysis")
                    
                    # Word cloud
                    st.write("#### Word Cloud")
                    if extracted_text.strip():
                        words = re.findall(r'\w+', extracted_text.lower()) # Check filtering
                        stop_words = set(STOPWORDS)
                        filtered_words = [w for w in words if w not in stop_words and len(w) > 1]
                        st.write("Filtered words:", filtered_words) 
                        
                        if filtered_words:  # Check if there are any words left after filtering
                            wordcloud_text = " ".join(filtered_words) 
                            # Generate word cloud
                            wc = WordCloud(width=800, height=400, 
                                         background_color='white', 
                                         colormap='viridis', 
                                         max_words=200)
                            wc.generate(" ".join(filtered_words))
                            
                            # Display
                            fig, ax = plt.subplots(figsize=(10, 5))
                            ax.imshow(wc, interpolation='bilinear')
                            ax.axis('off')
                            st.pyplot(fig)
                        else:
                            st.info("Not enough meaningful words to generate a word cloud. Try uploading a file with more content.")
                    else:
                         st.warning("No text content detected. Please upload a file with more words.")
                    
                    # Top words
                    st.write("#### Top Words")
                    stop_words = set(STOPWORDS)
                    filtered_words = [w for w in words if w not in stop_words and len(w) > 2]
                    
                    if filtered_words:
                        filtered_freq = Counter(filtered_words)
                        top_words = pd.DataFrame(filtered_freq.most_common(20), columns=['Word', 'Count'])
                        
                        fig = px.bar(top_words, x='Word', y='Count')
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.info("No significant words found for frequency analysis.")
                    
                    # Word length distribution
                    st.write("#### Word Length Distribution")
                    if filtered_words:
                        word_lengths = [len(w) for w in filtered_words]
                        word_length_freq = Counter(word_lengths)
                        length_df = pd.DataFrame(sorted(word_length_freq.items()), columns=['Length', 'Count'])
                        
                        fig = px.bar(length_df, x='Length', y='Count')
                        st.plotly_chart(fig, use_container_width=True)
                    else:
                        st.info("No significant words found for length analysis.")
                    
                    # Sentiment analysis
                    st.write("#### Sentiment Analysis")
                    if extracted_text.strip():
                        try:
                            from nltk.sentiment import SentimentIntensityAnalyzer
                            sia = SentimentIntensityAnalyzer()
                            sentiment = sia.polarity_scores(extracted_text)
                            
                            sentiment_df = pd.DataFrame({
                                'Metric': ['Negative', 'Neutral', 'Positive', 'Compound'],
                                'Value': [sentiment['neg'], sentiment['neu'], sentiment['pos'], sentiment['compound']]
                            })
                            
                            fig = px.bar(sentiment_df, x='Metric', y='Value', 
                                       color='Metric', 
                                       color_discrete_map={
                                           'Negative': 'red',
                                           'Neutral': 'gray',
                                           'Positive': 'green',
                                           'Compound': 'blue'
                                       })
                            st.plotly_chart(fig, use_container_width=True)
                        except:
                            st.info("Sentiment analysis requires NLTK with sentiment analysis models.")
                    else:
                        st.info("No text content for sentiment analysis.")
                    
                    # N-gram analysis
                    st.write("#### N-gram Analysis")
                    if filtered_words:
                        n_value = st.slider("Select n-gram size:", 2, 5, 2)
                        
                        # Generate n-grams
                        try:
                            from nltk.util import ngrams
                            n_grams = list(ngrams(filtered_words, n_value))
                            n_gram_freq = Counter([' '.join(g) for g in n_grams])
                            
                            if n_gram_freq:
                                n_gram_df = pd.DataFrame(n_gram_freq.most_common(15), 
                                                       columns=[f'{n_value}-gram', 'Count'])
                                
                                fig = px.bar(n_gram_df, x=f'{n_value}-gram', y='Count')
                                st.plotly_chart(fig, use_container_width=True)
                            else:
                                st.info(f"Not enough data to generate {n_value}-grams.")
                        except Exception as e:
                            st.error(f"Error in n-gram analysis: {str(e)}")
                    else:
                        st.info("Not enough text content for n-gram analysis.")

    # AI Chatbot Section
    st.subheader("💬 Chat with AI")
    user_input = st.text_area("Ask AI a question (Optional):")

    if st.button("Ask AI 🤖"):
        if extracted_text:
            prompt = f"Based on this file's content:\n{extracted_text[:3000]}\n\nAnswer the user's question: {user_input}"
        else:
            prompt = user_input  # If no file, just use user query

        response = generate_ai_content(prompt, st.session_state.api_key, st.session_state.api_model)
        st.success("🧠 AI Response:")
        st.write(response)
        
from iso639 import languages
from langdetect import detect
import pandas as pd
import docx
import json
from io import BytesIO

with tab5:
    st.header("🌍 AI-Powered Document Translator")
    
    # File uploader for document translation
    uploaded_file = st.file_uploader(
        "Upload a document (PDF, DOCX, TXT, CSV)", 
        type=["pdf", "docx", "txt", "csv"], 
        key="translator_file_uploader"
    )
    
    extracted_text = ""  # Store extracted text
    
    # Process the uploaded file and extract text
    if uploaded_file:
        extracted_text = extract_text_from_file(uploaded_file)  # Ensure this function is defined
        st.success("✅ File processed successfully!")
        st.text_area("Extracted Content:", extracted_text[:5000], height=200, key="translator_extracted_text")

    # Auto-detect source language
    detected_language = "Unknown"
    if extracted_text:
        try:
            detected_language = detect(extracted_text)
            detected_language_name = next((name for name, code in languages.items() if code == detected_language), "Unknown")
            st.info(f"🔍 Detected Source Language: **{detected_language_name} ({detected_language})**")
        except:
            st.warning("⚠️ Could not detect language. Please specify manually.")

    # Advanced translation options
    with st.expander("Advanced Translation Options"):
        translation_mode = st.selectbox(
            "Translation Mode:", 
            ["Full Translation", "Summary Translation", "Paraphrase Translation", "Technical Translation"], 
            key="translator_mode"
        )
        tone = st.selectbox(
            "Translation Tone:", 
            ["Formal", "Informal", "Neutral", "Friendly", "Academic"], 
            key="translator_tone"
        )
        output_format = st.selectbox(
            "Output Format:", 
            ["Plain Text", "Markdown", "HTML", "DOCX", "CSV", "JSON"], 
            key="translator_format"
        )
        max_chars = st.slider(
            "Max Characters to Translate:", 
            min_value=1000, max_value=20000, value=5000, step=500, 
            key="translator_max_chars"
        )
        parallel_translation = st.checkbox("Enable Parallel Translations (Compare Multiple Translations)", key="parallel_translation")
        pronunciation = st.checkbox("Include Pronunciation (For Language Learning)", key="pronunciation")

    # Fetch 200+ languages dynamically
    all_languages = {lang.name: lang.part1 for lang in languages if lang.part1}  # Get name & ISO code

    # Streamlit selectbox for choosing target language
    target_lang = st.selectbox(
        "Select Target Language:", 
        list(all_languages.keys()), 
        key="translator_language_select"
    )
    
    # Translation button and display translation
    if st.button("Translate 🌍", key="translator_translate_btn"):
        if extracted_text:
            lang_code = all_languages[target_lang]  # Convert display name to language code
            text_to_translate = extracted_text[:max_chars]  # Limit text length
            
            # Build an advanced prompt with options
            prompt = (
                f"Translation Mode: {translation_mode}\n"
                f"Tone: {tone}\n"
                f"Output Format: {output_format}\n"
                f"Pronunciation: {'Yes' if pronunciation else 'No'}\n"
                f"Translate the following text to {target_lang} ({lang_code}):\n\n{text_to_translate}"
            )
            
            translated_text = generate_ai_content(prompt, st.session_state.api_key, st.session_state.api_model)
            st.success("✅ Translation Complete:")
            st.write(translated_text)
            
            # Parallel Translations (if enabled)
            if parallel_translation:
                st.subheader("🔄 Parallel Translations")
                alternate_prompt = f"Provide an alternate translation for the text in {target_lang} ({lang_code}):\n\n{text_to_translate}"
                alternate_translation = generate_ai_content(alternate_prompt, st.session_state.api_key, st.session_state.api_model)
                
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("### 🎯 Primary Translation")
                    st.write(translated_text)
                with col2:
                    st.markdown("### 🛠 Alternate Translation")
                    st.write(alternate_translation)

            # Create downloadable DOCX file
            def create_docx(text):
                doc = docx.Document()
                doc.add_paragraph(text)
                doc_stream = BytesIO()
                doc.save(doc_stream)
                doc_stream.seek(0)  # Move to the beginning of the file
                return doc_stream
            
            # Create downloadable CSV file
            def create_csv(text):
                csv_data = pd.DataFrame([[target_lang, text]], columns=["Language", "Translation"])
                csv_stream = BytesIO()
                csv_data.to_csv(csv_stream, index=False)
                csv_stream.seek(0)
                return csv_stream

            # Export and copy options
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.download_button(
                    "Download as TXT", 
                    translated_text, 
                    file_name=f"translation_{target_lang}.txt"
                )
            with col2:
                docx_file = create_docx(translated_text)
                st.download_button(
                    "Download as DOCX", 
                    docx_file,  # ✅ Fixed: Pass binary stream instead of docx object
                    file_name=f"translation_{target_lang}.docx"
                )
            with col3:
                csv_file = create_csv(translated_text)
                st.download_button(
                    "Download as CSV", 
                    csv_file, 
                    file_name=f"translation_{target_lang}.csv"
                )
            with col4:
                st.download_button(
                    "Download as JSON", 
                    json.dumps({"language": target_lang, "text": translated_text}), 
                    file_name=f"translation_{target_lang}.json"
                )

        else:
            st.warning("⚠️ Please upload a document first!")


with tab6:
    st.header("⚡ AI Code Wizard")

    # Automatically select the AI model for coding tasks
    st.session_state.api_model = "gemini-2.5-pro-exp-03-25"

    # Choose AI task
    task = st.selectbox("What do you need help with?", 
                    ["Generate Code", "Debug Code", "Optimize Code", "Convert Code", "Explain Code",
                     "Add Comments", "Find Security Issues", "Write Unit Tests", "Generate API Documentation",
                     "Suggest Design Patterns", "Convert Pseudocode to Code", "Fix Compilation Errors",
                     "Analyze Code Performance"], 
                    key="code_task_radio")

    # Code Input (Text or File)
    code_input = st.text_area("Enter your code or request:", height=200, key="code_input")

    uploaded_code = st.file_uploader("Upload a code file", type=["py", "js", "cpp", "java", "vhdl"], key="code_file_uploader")
    
    if uploaded_code:
        code_input = uploaded_code.read().decode("utf-8")
        st.text_area("Uploaded Code:", code_input, height=200, key="uploaded_code_display")

    # AI Action Button
    if st.button("✨ Magic Code AI", key="code_magic_button"):
        if code_input.strip():
            if task == "Generate Code":
                prompt = f"Write clean, efficient code for: {code_input}"
            elif task == "Debug Code":
                prompt = f"Find and fix errors in this code:\n\n{code_input}"
            elif task == "Optimize Code":
                prompt = f"Refactor and optimize this code to improve efficiency:\n\n{code_input}"
            elif task == "Convert Code":
                prompt = f"Convert this code to another programming language:\n\n{code_input}"
            elif task == "Explain Code":
                prompt = f"Explain what this code does in simple terms:\n\n{code_input}"
            elif task == "Add Comments":
                prompt = f"Add detailed comments to this code for better understanding:\n\n{code_input}"
            elif task == "Find Security Issues":
                prompt = f"Analyze this code and highlight security vulnerabilities:\n\n{code_input}"
            elif task == "Write Unit Tests":
                prompt = f"Generate unit tests for this code:\n\n{code_input}"
            elif task == "Generate API Documentation":
                prompt = f"Create API documentation for this code:\n\n{code_input}"
            elif task == "Suggest Design Patterns":
                prompt = f"Suggest an appropriate design pattern for this code and explain why:\n\n{code_input}"
            elif task == "Convert Pseudocode to Code":
                prompt = f"Convert this pseudocode into actual code:\n\n{code_input}"
            elif task == "Fix Compilation Errors":
                prompt = f"Fix the compilation errors in this code:\n\n{code_input}"
            elif task == "Analyze Code Performance":
                prompt = f"Analyze the performance of this code and suggest improvements:\n\n{code_input}"

            # AI Code Processing
            ai_response = generate_ai_content(prompt, st.session_state.api_key, st.session_state.api_model)

            st.success("✨ AI Code Response:")
            st.code(ai_response, language="python")  # Adjust language based on task
        else:
            st.warning("⚠️ Please enter some code or upload a file.")

# Content generation section
st.header("✨ Create Content")

# Display selected tool or default
selected_tool = st.session_state.get('selected_tool', 'Smart Content Creator')
st.markdown(f"### Currently using: **{selected_tool}**")

# Content prompt area
user_prompt = st.text_area("Enter your prompt:", "", height=150)

# Optional file uploader (Excluding images and music files)
uploaded_file = st.file_uploader("Upload a text-based file (PDF, DOCX, TXT, CSV)", type=["pdf", "docx", "txt", "csv"])

# Function to extract text from uploaded file
if uploaded_file:
    extracted_text = extract_text_from_file(uploaded_file)  # Ensure this function is defined in your code
    if extracted_text:
        st.success("✅ File uploaded successfully!")
        st.text_area("Extracted Content:", extracted_text[:5000], height=200, key="uploaded_content_area")
        
        # Option to merge extracted text with user input
        if st.checkbox("Include extracted content in prompt"):
            user_prompt += f"\n\n{extracted_text}"

# Store updated prompt in session state
st.session_state['user_prompt'] = user_prompt


import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import numpy as np
from sklearn.ensemble import IsolationForest
from sklearn.feature_selection import mutual_info_regression, mutual_info_classif


import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import numpy as np
from sklearn.ensemble import IsolationForest
from sklearn.feature_selection import mutual_info_regression, mutual_info_classif

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import numpy as np
from sklearn.ensemble import IsolationForest
from sklearn.feature_selection import mutual_info_regression, mutual_info_classif

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import numpy as np
from sklearn.ensemble import IsolationForest
from sklearn.feature_selection import mutual_info_regression, mutual_info_classif
import plotly.express as px
import scipy.cluster.hierarchy as sch
from sklearn.preprocessing import StandardScaler
from wordcloud import WordCloud

with tab7:
    st.header("📊 Data Visualization & Insights")

    # File uploader for CSV and data files
    uploaded_file = st.file_uploader("Upload a dataset (CSV, Excel)", type=["csv", "xlsx"])

    if uploaded_file:
        # Read data
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)

        # Display basic information
        st.subheader("Dataset Overview")
        st.write(df.head())  # Show first few rows

        # Show basic statistics
        st.subheader("Basic Insights")
        st.write(df.describe())

        # Select numeric and categorical columns
        numeric_columns = df.select_dtypes(include=["number"]).columns.tolist()
        categorical_columns = df.select_dtypes(include=["object"]).columns.tolist()

        if numeric_columns:
            st.subheader("📈 Choose a Visualization")

            # Select visualization type
            plot_type = st.selectbox(
                "Select a plot to display",
                [
                    "Pairplot", "Correlation Heatmap", "Scatter Plot", "Histogram", 
                    "Boxplot", "Violin Plot", "Line Plot", "Bar Chart", "KDE Plot", 
                    "Pie Chart", "Swarm Plot", "Stacked Area Chart", "Parallel Coordinates",
                    "Dendrogram", "Radar Chart", "3D Scatter Plot", "Joint Plot", 
                    "Autocorrelation Plot", "Treemap", "Word Cloud", "Anomaly Detection", 
                    "Feature Importance"
                ]
            )

            # Pairplot
            if plot_type == "Pairplot":
                selected_cols = st.multiselect("Select columns for Pairplot", numeric_columns, default=numeric_columns[:3])
                if selected_cols:
                    fig = sns.pairplot(df[selected_cols])
                    st.pyplot(fig)

            # Correlation Heatmap
            elif plot_type == "Correlation Heatmap":
                selected_cols = st.multiselect("Select columns for Heatmap", numeric_columns, default=numeric_columns[:5])
                if selected_cols:
                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.heatmap(df[selected_cols].corr(), annot=True, cmap="coolwarm", ax=ax)
                    st.pyplot(fig)

            # Scatter Plot
            elif plot_type == "Scatter Plot":
                x_axis = st.selectbox("Select X-axis", numeric_columns)
                y_axis = st.selectbox("Select Y-axis", numeric_columns)
                fig = px.scatter(df, x=x_axis, y=y_axis)
                st.plotly_chart(fig)

            # Histogram
            elif plot_type == "Histogram":
                selected_cols = st.multiselect("Select multiple columns for Histogram", numeric_columns, default=[numeric_columns[0]])
                for col in selected_cols:
                    fig, ax = plt.subplots()
                    sns.histplot(df[col], bins=30, kde=True, ax=ax)
                    st.pyplot(fig)

            # Boxplot
            elif plot_type == "Boxplot":
                selected_cols = st.multiselect("Select multiple columns for Boxplot", numeric_columns, default=[numeric_columns[0]])
                for col in selected_cols:
                    fig, ax = plt.subplots()
                    sns.boxplot(y=df[col], ax=ax)
                    st.pyplot(fig)

            # Violin Plot
            elif plot_type == "Violin Plot":
                selected_cols = st.multiselect("Select multiple columns for Violin Plot", numeric_columns, default=[numeric_columns[0]])
                for col in selected_cols:
                    fig, ax = plt.subplots()
                    sns.violinplot(y=df[col], ax=ax)
                    st.pyplot(fig)

            # Line Plot
            elif plot_type == "Line Plot":
                selected_cols = st.multiselect("Select multiple columns for Line Plot", numeric_columns, default=[numeric_columns[0]])
                for col in selected_cols:
                    fig, ax = plt.subplots()
                    sns.lineplot(x=df.index, y=df[col], ax=ax)
                    st.pyplot(fig)

            # Bar Chart
            elif plot_type == "Bar Chart":
                if categorical_columns:
                    cat_col = st.selectbox("Select a categorical column", categorical_columns)
                    num_cols = st.multiselect("Select multiple numeric columns", numeric_columns, default=[numeric_columns[0]])
                    for col in num_cols:
                        fig, ax = plt.subplots()
                        df.groupby(cat_col)[col].mean().plot(kind="bar", ax=ax)
                        st.pyplot(fig)

            # KDE Plot
            elif plot_type == "KDE Plot":
                selected_cols = st.multiselect("Select multiple columns for KDE Plot", numeric_columns, default=[numeric_columns[0]])
                for col in selected_cols:
                    fig, ax = plt.subplots()
                    sns.kdeplot(df[col], fill=True, ax=ax)
                    st.pyplot(fig)

            # Pie Chart
            elif plot_type == "Pie Chart":
                if categorical_columns:
                    selected_cols = st.multiselect("Select categorical columns for Pie Chart", categorical_columns, default=[categorical_columns[0]])
                    for col in selected_cols:
                        fig, ax = plt.subplots()
                        df[col].value_counts().plot(kind="pie", autopct="%1.1f%%", ax=ax)
                        st.pyplot(fig)
            
            # Swarm Plot
            elif plot_type == "Swarm Plot":
                if categorical_columns:
                    cat_col = st.selectbox("Select a categorical column", categorical_columns)
                    num_col = st.selectbox("Select a numeric column", numeric_columns)
                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.swarmplot(x=df[cat_col], y=df[num_col], ax=ax)
                    plt.xticks(rotation=45)
                    st.pyplot(fig)
                else:
                    st.warning("Swarm Plot requires at least one categorical column.")

            # Stacked Area Chart
            elif plot_type == "Stacked Area Chart":
                selected_cols = st.multiselect("Select multiple columns for Stacked Area Chart", numeric_columns, default=numeric_columns[:3])
                if selected_cols:
                    fig = px.area(df, x=df.index, y=selected_cols)
                    st.plotly_chart(fig)

            # Parallel Coordinates
            elif plot_type == "Parallel Coordinates":
                selected_cols = st.multiselect("Select columns for Parallel Coordinates", numeric_columns, default=numeric_columns[:5])
                if selected_cols and len(selected_cols) >= 3:
                    color_col = st.selectbox("Select column for color", selected_cols)
                    fig = px.parallel_coordinates(df, dimensions=selected_cols, color=color_col)
                    st.plotly_chart(fig)
                else:
                    st.warning("Parallel Coordinates requires at least 3 numeric columns.")

            # Dendrogram
            elif plot_type == "Dendrogram":
                selected_cols = st.multiselect("Select columns for Dendrogram", numeric_columns, default=numeric_columns[:5])
                if selected_cols:
                    fig, ax = plt.subplots(figsize=(10, 8))
                    Z = hierarchy.linkage(df[selected_cols], 'ward')
                    dn = hierarchy.dendrogram(Z, ax=ax)
                    st.pyplot(fig)

            # Radar Chart
            elif plot_type == "Radar Chart":
                selected_cols = st.multiselect("Select columns for Radar Chart", numeric_columns, default=numeric_columns[:5])
                if selected_cols and len(selected_cols) >= 3:
                    df_radar = df[selected_cols].copy()
                    for col in selected_cols:
                        df_radar[col] = (df_radar[col] - df_radar[col].min()) / (df_radar[col].max() - df_radar[col].min())
                    samples = min(5, len(df_radar))
                    sample_rows = df_radar.sample(samples).index
                    fig = go.Figure()
                    for i, row in enumerate(sample_rows):
                        fig.add_trace(go.Scatterpolar(r=df_radar.loc[row].values, theta=selected_cols, fill='toself', name=f'Sample {i+1}'))
                    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 1])))
                    st.plotly_chart(fig)
                else:
                    st.warning("Radar Chart requires at least 3 numeric columns.")

            # 3D Scatter Plot
            elif plot_type == "3D Scatter Plot":
                if len(numeric_columns) >= 3:
                    x_axis = st.selectbox("Select X-axis", numeric_columns, index=0)
                    y_axis = st.selectbox("Select Y-axis", numeric_columns, index=min(1, len(numeric_columns)-1))
                    z_axis = st.selectbox("Select Z-axis", numeric_columns, index=min(2, len(numeric_columns)-1))
                    fig = px.scatter_3d(df, x=x_axis, y=y_axis, z=z_axis)
                    st.plotly_chart(fig)
                else:
                    st.warning("3D Scatter Plot requires at least 3 numeric columns.")

            # Joint Plot
            elif plot_type == "Joint Plot":
                x_col = st.selectbox("Select X-axis column", numeric_columns, index=0)
                y_col = st.selectbox("Select Y-axis column", numeric_columns, index=min(1, len(numeric_columns)-1))
                fig = sns.jointplot(x=df[x_col], y=df[y_col], kind="hex")
                st.pyplot(fig)

            # Autocorrelation Plot
            elif plot_type == "Autocorrelation Plot":
                selected_col = st.selectbox("Select column for Autocorrelation", numeric_columns)
                fig, ax = plt.subplots(figsize=(10, 6))
                pd.plotting.autocorrelation_plot(df[selected_col], ax=ax)
                st.pyplot(fig)

            # Treemap
            elif plot_type == "Treemap":
                if categorical_columns:
                    cat_cols = st.multiselect("Select categorical columns (hierarchy levels)", categorical_columns, default=[categorical_columns[0]])
                    val_col = st.selectbox("Select value column", numeric_columns)
                    if cat_cols and val_col:
                        fig = px.treemap(df, path=cat_cols, values=val_col)
                        st.plotly_chart(fig)
                else:
                    st.warning("Treemap requires at least one categorical column.")

            # Word Cloud
            elif plot_type == "Word Cloud":
                if categorical_columns:
                    text_col = st.selectbox("Select text column for Word Cloud", categorical_columns)
                    stop_words = st.text_input("Enter additional stop words (comma separated)")
                    stop_words_list = [word.strip() for word in stop_words.split(',')] if stop_words else []
                    
                    # Combine all text
                    text = ' '.join(df[text_col].dropna().astype(str))
                    
                    # Generate word cloud
                    wordcloud = WordCloud(width=800, height=400, background_color='white', stopwords=set(STOPWORDS).union(stop_words_list)).generate(text)
                    
                    # Display
                    fig, ax = plt.subplots(figsize=(10, 5))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis('off')
                    st.pyplot(fig)
                else:
                    st.warning("Word Cloud requires at least one text column.")

            # Anomaly Detection
            elif plot_type == "Anomaly Detection":
                selected_cols = st.multiselect("Select columns for anomaly detection", numeric_columns, default=[numeric_columns[0]])
                contamination = st.slider("Contamination percentage", 0.01, 0.1, 0.05)
                if selected_cols:
                    model = IsolationForest(contamination=contamination, random_state=42)
                    df["Anomaly"] = model.fit_predict(df[selected_cols])
                    outliers = df[df["Anomaly"] == -1]
                    st.write(f"Detected {len(outliers)} potential anomalies.")

                    fig, ax = plt.subplots()
                    sns.scatterplot(x=df.index, y=df[selected_cols[0]], hue=df["Anomaly"], palette={1: "blue", -1: "red"}, ax=ax)
                    st.pyplot(fig)

            # Feature Importance
            elif plot_type == "Feature Importance":
                if categorical_columns:
                    target_col = st.selectbox("Select target column", categorical_columns)
                    df[target_col] = pd.to_datetime(df[target_col], errors="coerce").astype(int) / 10**9
                    is_classification = df[target_col].nunique() < 10
                    importance = mutual_info_classif(df[numeric_columns], df[target_col]) if is_classification else mutual_info_regression(df[numeric_columns], df[target_col])
                    feature_importance = pd.DataFrame({"Feature": numeric_columns, "Importance": importance}).sort_values(by="Importance", ascending=False)
                    st.write(feature_importance)
                    fig = px.bar(feature_importance, x="Importance", y="Feature", orientation="h")
                    st.plotly_chart(fig)

        else:
            st.warning("No numerical columns available for visualization.")
    else:
        st.info("Upload a dataset to generate insights and visualizations.")



with tab8:
    st.header("🎤 AI-Powered Mock Interviewer")
    st.markdown("### Prepare for your next job interview with AI-generated questions.")

    # Job Role Selection
    col1, col2 = st.columns(2)
    with col1:
        job_role = st.text_input("Enter Job Role (e.g., Data Scientist, Software Engineer):")
    with col2:
        experience_level = st.selectbox("Experience Level:", ["Entry Level", "Mid-Level", "Senior-Level", "Executive"])

    # Interview Type Selection
    interview_type = st.radio("Interview Type:", ["Technical", "Behavioral", "Case Study", "HR"], horizontal=True)

    # Advanced Options
    with st.expander("🎯 Customize Your Mock Interview"):
        col1, col2 = st.columns(2)
        with col1:
            difficulty = st.slider("Difficulty Level:", 1, 10, 5)
        with col2:
            question_count = st.slider("Number of Questions:", 3, 20, 5)

    # Generate Questions
    if st.button("Start Mock Interview 🎙️"):
        interview_prompt = f"""
        Conduct a {interview_type} interview for a {experience_level} {job_role}. 
        Ask {question_count} questions with increasing difficulty from level {difficulty}.
        """

        with st.spinner("Generating interview questions..."):
            interview_questions = generate_ai_content(interview_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("🎤 AI-Generated Mock Interview:")
        st.write(interview_questions)

        # Download Options
        st.download_button("📥 Download Questions", interview_questions, "mock_interview.txt")


with tab9:
    st.header("📧 AI-Powered Email Assistant")
    st.markdown("### Generate, summarize, and rewrite emails with AI.")

    # Email Input
    email_content = st.text_area("✍️ Enter Email Content or Brief:")

    # Tone Selection
    tone = st.radio("Select Tone:", ["Formal", "Informal", "Neutral", "Persuasive", "Apologetic"])

    # Action Selection
    action = st.radio("What would you like to do?", 
                      ["Generate Email", "Summarize Email", "Rewrite Email", "Quick Reply", "Follow-Up Suggestion"], 
                      horizontal=True)

    # Email Personalization
    recipient_type = st.selectbox("Recipient Type:", ["Boss", "Colleague", "Client", "Friend", "General"])

    # Attachment Upload (for Email Summarization)
    uploaded_file = st.file_uploader("📎 Upload a file (PDF, DOCX, TXT) for summary (Optional)", 
                                     type=["pdf", "docx", "txt"])

    # Email Scheduling Assistant
    urgency = st.radio("📅 Email Urgency:", ["Immediate", "Within 24 Hours", "End of Week", "Next Week"])
    if urgency == "Immediate":
        best_time = "Send Now 🚀"
    elif urgency == "Within 24 Hours":
        best_time = "Send within the next few hours ⏳"
    elif urgency == "End of Week":
        best_time = "Send by Friday afternoon 📆"
    else:
        best_time = "Schedule for next Monday ⏰"
    
    st.markdown(f"**🕒 Suggested Send Time: {best_time}**")

    # Generate Button
    if st.button("✉️ Process Email"):
        email_prompt = f"""
        Perform '{action}' on this email: {email_content} with a '{tone}' tone.
        Tailor it for a '{recipient_type}' and consider urgency level '{urgency}'.
        """

        # Handle File Upload for Summary
        if uploaded_file is not None:
            file_text = extract_text_from_file(uploaded_file)
            email_prompt += f"\n\n[Attachment Summary: {file_text}]"

        with st.spinner("Processing..."):
            processed_email = generate_ai_content(email_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Email:")
        st.write(processed_email)

        # Download Button
        st.download_button("📥 Download Email", processed_email, "email.txt")

with tab10:
    st.header("📊 AI-Powered Spreadsheet Formula Generator")
    st.markdown("### Convert plain English into powerful spreadsheet formulas.")

    # User Input for Natural Language Query
    query = st.text_area("📝 Describe what you need (e.g., 'Extract domain from email'):")

    # Syntax Selection
    syntax = st.radio("Choose Formula Syntax:", ["Excel", "Google Sheets", "LibreOffice Calc"])

    # Formula Type Selection
    formula_type = st.radio("Formula Type:", ["Single Cell", "Multi-Cell/Range", "Array Formula", "SQL Query"], horizontal=True)

    # Debug & Optimization Option
    debug_formula = st.checkbox("Enable Formula Debugging & Optimization", value=True)

    # Function Recommendations
    function_suggestions = st.checkbox("Suggest Best Functions for My Task", value=True)

    # Data Cleaning & Formatting
    data_cleaning = st.checkbox("Enable AI Data Cleaning & Auto-Formatting", value=True)

    # Auto-Generate Pivot Tables & Charts
    generate_pivot = st.checkbox("Generate Pivot Tables & Charts Automatically", value=False)

    # Bulk Formula Generation for Large Datasets
    bulk_formula = st.checkbox("Apply Formula in Bulk for Large Datasets", value=False)

    # AI-Powered Conditional Formatting
    conditional_formatting = st.checkbox("Enable Smart Conditional Formatting", value=False)

    # Generate Button
    if st.button("🔢 Generate Formula"):
        formula_prompt = f"""
        Convert this request into a {formula_type} formula for {syntax}: {query}.
        {"Provide a breakdown and explanation." if debug_formula else ""}
        {"Suggest the best spreadsheet functions for this task." if function_suggestions else ""}
        {"Optimize and debug the formula for efficiency." if debug_formula else ""}
        {"Ensure the formula works well with large datasets." if bulk_formula else ""}
        {"Automatically clean and format the data before applying the formula." if data_cleaning else ""}
        {"Suggest conditional formatting rules if applicable." if conditional_formatting else ""}
        {"If relevant, generate a Pivot Table or Chart based on the query." if generate_pivot else ""}
        """

        with st.spinner("Generating formula..."):
            generated_formula = generate_ai_content(formula_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Formula:")
        st.code(generated_formula, language="plaintext")

        # Download Button
        st.download_button("📥 Download Formula", generated_formula, "formula.txt")



with tab11:
    st.header("🎙️ AI-Based Podcast Script Generator")
    st.markdown("### Create engaging, professional, and unique podcast scripts!")

    # Podcast Topic & Duration
    topic = st.text_input("🎯 Enter Podcast Topic:")
    duration = st.slider("⏳ Select Podcast Duration (minutes):", 5, 90, 20)
    
    # Podcast Style & Tone
    style = st.radio("🎭 Choose Style:", ["Conversational", "Storytelling", "Humorous", "Data-Driven", "Inspirational"])
    tone = st.radio("🗣️ Select Tone:", ["Casual", "Professional", "Energetic", "Serious"], horizontal=True)

    # Word Count Control
    col1, col2 = st.columns(2)
    with col1:
        min_words = st.number_input("🔡 Min Words:", min_value=100, max_value=5000, value=300, step=50)
    with col2:
        max_words = st.number_input("🔠 Max Words:", min_value=200, max_value=10000, value=1500, step=100)

    # Script Detail Level
    script_detail = st.select_slider("📜 Script Detail Level:", 
                                     options=["Short", "Standard", "Comprehensive"], 
                                     value="Standard")

    # Podcast Features & Enhancements
    add_intro_outro = st.checkbox("🎶 Auto-Generated Podcast Intro & Outro", value=True)
    simulate_cohost_guest = st.checkbox("🧑‍🤝‍🧑 Simulate Multi-Host & Guest Interactions", value=True)
    add_trending_news = st.checkbox("🌎 Include Real-Time News & Trends", value=False)
    add_music_sfx = st.checkbox("🎵 Suggest Background Music & Sound Effects", value=False)
    ai_character_speaking = st.checkbox("🤖 Generate Script in a Famous Figure's Style", value=False)
    add_hot_takes = st.checkbox("💡 Include Debates & Controversial Takes", value=False)
    future_predictions = st.checkbox("🔮 AI-Powered Future Insights & Trends", value=False)
    generate_ads = st.checkbox("📜 Create Sponsor Ad Scripts", value=False)
    auto_schedule = st.checkbox("🛠️ Auto-Generate Social Media Promo Posts", value=False)
    multi_language = st.checkbox("📢 Translate to Multiple Languages", value=False)

    # Generate Button
    if st.button("📜 Generate Podcast Script"):
        podcast_prompt = f"""
        Generate a {script_detail} {duration}-minute podcast script on '{topic}' in {style} style with a {tone} tone.
        Ensure the script is between {min_words} and {max_words} words.
        {"Include a dynamic podcast intro & outro." if add_intro_outro else ""}
        {"Simulate engaging conversations between a host and a guest." if simulate_cohost_guest else ""}
        {"Incorporate real-world trending news & insights." if add_trending_news else ""}
        {"Suggest background music or sound effects for key moments." if add_music_sfx else ""}
        {"Write the script in the style of a well-known personality or character." if ai_character_speaking else ""}
        {"Add a controversial debate section or counterpoints for engagement." if add_hot_takes else ""}
        {"Provide AI-driven future predictions related to the topic." if future_predictions else ""}
        {"Generate sponsor messages relevant to the topic." if generate_ads else ""}
        {"Auto-generate social media post copy for promoting this episode." if auto_schedule else ""}
        {"Translate the script into multiple languages with regional adaptation." if multi_language else ""}
        """

        with st.spinner("Generating podcast script..."):
            podcast_script = generate_ai_content(podcast_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Podcast Script:")
        st.text_area("🎙️ Your Podcast Script:", podcast_script, height=300)

        # Download Button
        st.download_button("📥 Download Script", podcast_script, "podcast_script.txt")


with tab12:
    st.header("🎯 AI-Powered Learning Path Generator")
    st.markdown("### Get a personalized learning roadmap based on your skills and goals!")

    # User Inputs
    current_skill = st.text_input("🛠️ Current Skills (e.g., Python, Data Science, Web Dev):")
    goal = st.text_input("🎯 Learning Goal (e.g., Become a Machine Learning Engineer):")

    # Select Learning Style
    learning_style = st.radio("📚 Preferred Learning Style:", 
                                  ["Video-Based", "Hands-On Projects", "Reading-Focused", "Hybrid"])

    # Select Duration
    duration = st.slider("⏳ Select Learning Duration (Weeks):", 1, 52, 12)

    # Skill Level Selection
    skill_level = st.radio("📊 Current Proficiency Level:", 
                           ["Beginner", "Intermediate", "Advanced"], horizontal=True)

    # Word Count Control
    col1, col2 = st.columns(2)
    with col1:
        min_words = st.number_input("🔡 Min Words:", min_value=100, max_value=5000, value=500, step=50,key="min_words_input")
    with col2:
        max_words = st.number_input("🔠 Max Words:", min_value=200, max_value=10000, value=1500, step=100,key="max_words_input")

    # Additional Features
    ai_gap_analysis = st.checkbox("🛠️ AI Skill Gap Analysis", value=True)
    personalized_challenges = st.checkbox("📌 Generate Weekly Challenges", value=True)
    certification_mapping = st.checkbox("🎓 Recommend Certifications & Courses", value=True)
    resume_boost = st.checkbox("📂 AI Resume Enhancement", value=False)
    adaptive_learning = st.checkbox("🧠 Adaptive Learning (Dynamic Updates)", value=False)
    time_estimate = st.checkbox("⏳ Time Commitment Estimator", value=True)
    mentor_suggestions = st.checkbox("👥 Suggest Mentors & Learning Communities", value=False)
    quizzes = st.checkbox("📝 Generate Self-Assessment Quizzes", value=True)
    job_insights = st.checkbox("📢 Industry Trends & Job Market Insights", value=True)

    # Generate Learning Path
    if st.button("📜 Generate Learning Path"):
        learning_prompt = f"""
        Create a {duration}-week personalized learning roadmap for someone with {skill_level} skills in {current_skill} 
        who wants to achieve the goal: '{goal}'.
        Preferred learning style: {learning_style}.
        Ensure the response is between {min_words} and {max_words} words.
        {"Analyze the skill gap between their current skills and their goal." if ai_gap_analysis else ""}
        {"Generate weekly challenges to test their learning progress." if personalized_challenges else ""}
        {"Recommend certifications, books, and online courses." if certification_mapping else ""}
        {"Suggest how to add these skills to a resume for better job opportunities." if resume_boost else ""}
        {"Provide an adaptive learning roadmap that changes based on progress." if adaptive_learning else ""}
        {"Estimate the required daily/weekly study time to reach the goal." if time_estimate else ""}
        {"Suggest relevant mentorship programs and online learning communities." if mentor_suggestions else ""}
        {"Include self-assessment quizzes to measure progress." if quizzes else ""}
        {"Provide industry trends and job market insights for this skill." if job_insights else ""}
        """

        with st.spinner("Generating your learning path..."):
            learning_path = generate_ai_content(learning_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Learning Path:")
        st.text_area("🎯 Your Personalized Learning Roadmap:", learning_path, height=300)

        # Download Option
        st.download_button("📥 Download Learning Path", learning_path, "learning_path.txt")



with tab13:
    st.header("📝 AI-Powered Meeting Minutes Generator")
    st.markdown("### Automatically generate structured meeting summaries from transcripts or audio files.")

    # Upload Meeting File or Paste Custom Transcript
    uploaded_meeting = st.file_uploader("📂 Upload Meeting Transcript or Audio (TXT, DOCX, PDF, MP3, WAV)", 
                                        type=["txt", "docx", "pdf", "mp3", "wav"])

    st.markdown("**OR**")
    custom_transcript = st.text_area("✍️ Paste Your Meeting Transcript (Optional):", height=200)

    # Key Features Selection
    extract_decisions = st.checkbox("🎯 Extract Key Decisions", value=True)
    extract_action_items = st.checkbox("✅ Identify Action Items & Owners", value=True)
    extract_deadlines = st.checkbox("📅 Highlight Deadlines", value=True)
    sentiment_analysis = st.checkbox("💡 Perform Sentiment Analysis", value=False)
    search_keywords = st.text_input("🔎 Search for Specific Keywords (Optional):")

    # Process Meeting File or Custom Input
    extracted_meeting_text = ""
    if uploaded_meeting is not None:
        extracted_meeting_text = extract_text_from_file(uploaded_meeting)  # Function to extract text from uploaded file
        st.success("✅ Meeting file uploaded successfully! AI will analyze it.")
    elif custom_transcript:
        extracted_meeting_text = custom_transcript
        st.success("✅ Custom transcript added! AI will analyze it.")

    # Generate Meeting Summary
    if extracted_meeting_text and st.button("📜 Generate Meeting Minutes"):
        meeting_prompt = f"""
        Analyze the following meeting transcript and generate structured minutes.
        {"Extract key decisions made." if extract_decisions else ""}
        {"Identify action items, assigned members, and deadlines." if extract_action_items else ""}
        {"Highlight upcoming deadlines and due dates." if extract_deadlines else ""}
        {"Perform sentiment analysis to detect agreement, disagreement, or conflicts." if sentiment_analysis else ""}
        {"Include a search-based summary for the keyword(s): " + search_keywords if search_keywords else ""}
        
        Meeting Transcript:
        {extracted_meeting_text}
        """

        with st.spinner("Processing meeting minutes..."):
            meeting_summary = generate_ai_content(meeting_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Meeting Minutes:")
        st.text_area("📜 Your Meeting Summary:", meeting_summary, height=300)

        # Download & Copy Options
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("📥 Download Minutes", meeting_summary, "meeting_minutes.txt")
        with col2:
            st.button("📋 Copy to Clipboard", on_click=lambda: st.write(
                "<script>navigator.clipboard.writeText(`" + meeting_summary.replace("`", "\\`") + "`);</script>", 
                unsafe_allow_html=True))



with tab14:
    st.header("🚀 AI-Powered Startup Idea Validator")
    st.markdown("### Validate your business idea with AI-driven market analysis!")

    # Startup Idea Input
    startup_idea = st.text_area("📝 Describe Your Startup Idea (e.g., 'AI-Powered Resume Builder'):")
    
    # Industry Selection
    industry = st.radio("🏢 Select Industry:", 
                            ["Technology", "Healthcare", "Finance", "E-Commerce", "Education", "Other"])
    
    # Business Model Selection
    business_model = st.radio("💰 Select Revenue Model:", 
                              ["Subscription-Based", "Ad-Supported", "Freemium", "E-Commerce", "Enterprise Sales"], horizontal=True)

    # Additional AI Analysis Features
    competitor_analysis = st.checkbox("🏆 Include Competitor Research", value=True)
    market_potential = st.checkbox("📊 Analyze Market Demand", value=True)
    scalability_risks = st.checkbox("🚀 Assess Scalability & Risks", value=True)
    future_trends = st.checkbox("🔮 Predict Future Industry Trends", value=True)
    branding_advice = st.checkbox("📢 Provide Marketing & Branding Strategies", value=True)
    swot_analysis = st.checkbox("🧠 Generate SWOT Analysis", value=True)
    target_audience = st.checkbox("📍 Identify Target Audience", value=True)
    usp_finder = st.checkbox("💡 Find Unique Selling Proposition (USP)", value=True)
    mvp_plan = st.checkbox("🏗️ Create Minimum Viable Product (MVP) Plan", value=True)
    growth_strategy = st.checkbox("📈 Suggest Growth & Scaling Strategy", value=True)
    funding_readiness = st.checkbox("💵 Evaluate Funding & Investor Readiness", value=True)
    go_to_market = st.checkbox("🎯 Develop Go-To-Market Strategy", value=True)
    global_vs_local = st.checkbox("🌍 Recommend Global vs Local Expansion", value=True)
    tech_stack = st.checkbox("🛠️ Suggest Best Tech Stack & Tools", value=True)
    business_plan = st.checkbox("📜 Generate Business Plan Summary", value=True)

    # Generate Startup Report
    if st.button("📈 Validate My Startup Idea"):
        startup_prompt = f"""
        Analyze the feasibility of the startup idea: '{startup_idea}' in the {industry} industry.
        Revenue Model: {business_model}.
        {"Include competitor research." if competitor_analysis else ""}
        {"Analyze the market potential and demand." if market_potential else ""}
        {"Assess scalability challenges and risks." if scalability_risks else ""}
        {"Predict future trends in this industry." if future_trends else ""}
        {"Provide branding and marketing strategies." if branding_advice else ""}
        {"Generate a SWOT analysis (Strengths, Weaknesses, Opportunities, Threats)." if swot_analysis else ""}
        {"Identify the ideal target audience and customer personas." if target_audience else ""}
        {"Find a Unique Selling Proposition (USP) to differentiate the startup." if usp_finder else ""}
        {"Create a Minimum Viable Product (MVP) development plan." if mvp_plan else ""}
        {"Suggest growth, scaling, and funding strategies." if growth_strategy else ""}
        {"Assess funding readiness and provide investor pitching advice." if funding_readiness else ""}
        {"Develop a go-to-market strategy and customer acquisition plan." if go_to_market else ""}
        {"Evaluate whether this startup should focus on a local or global market." if global_vs_local else ""}
        {"Suggest the best tech stack and AI tools for development." if tech_stack else ""}
        {"Generate a structured business plan summary for investors." if business_plan else ""}
        """

        with st.spinner("Analyzing your startup idea..."):
            startup_report = generate_ai_content(startup_prompt, st.session_state.api_key, st.session_state.api_model)

        st.success("✅ AI-Generated Startup Validation Report:")
        st.text_area("📊 Your Startup Analysis:", startup_report, height=300)

        # Download & Copy Options
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("📥 Download Report", startup_report, "startup_report.txt")
        with col2:
            st.button("📋 Copy to Clipboard", on_click=lambda: st.write(
                "<script>navigator.clipboard.writeText(`" + startup_report.replace("`", "\\`") + "`);</script>", 
                unsafe_allow_html=True))


with tab15:
    st.header("🔮 AI Search & Knowledge Assistant")
    st.markdown("### Query documents, websites, or GitHub repos with AI-powered search!")
    
    # Select search type
    search_type = st.radio("📚 What would you like to search?", 
                          ["Website Content", "GitHub Repository", "Document (PDF/TXT)", "Deep Research"],
                          horizontal=True)
    
    if search_type == "Website Content":
        st.markdown("### 🌐 Website QA Agent")
        website_url = st.text_input("🔗 Enter website URL to analyze (e.g., https://example.com):")
        website_depth = st.slider("🕸️ Crawling Depth", min_value=1, max_value=5, value=2, 
                                help="Higher depth means more pages will be crawled")
        website_query = st.text_area("❓ What would you like to know about this website?")
        
    elif search_type == "GitHub Repository":
        st.markdown("### 🐙 GitHub Repo QA")
        repo_url = st.text_input("🔗 Enter GitHub repo URL (e.g., https://github.com/username/repo):")
        include_options = st.multiselect("📂 Select what to include:", 
                                        ["README", "Code Files", "Issues", "Pull Requests", "Discussions"],
                                        default=["README", "Code Files"])
        file_types = st.text_input("🔠 File extensions to include (comma-separated, e.g., py,md,js):", value="py,md,js,html,css")
        repo_query = st.text_area("❓ What would you like to know about this repository?")
        
    elif search_type == "Deep Research":
        st.markdown("### 🧠 Deep Research")
        deep_research_topic = st.text_area("📝 Enter your research topic:", 
                                        help="AI will generate comprehensive research on this topic")
        
    else:  # Document search
        st.markdown("### 📄 Document QA")
        uploaded_file = st.file_uploader("📎 Upload a document (PDF, TXT, DOCX):", type=["pdf", "txt", "docx"])
        if uploaded_file:
            st.success(f"✅ File '{uploaded_file.name}' uploaded successfully")
        doc_query = st.text_area("❓ What would you like to know about this document?")
    
    # Common settings
    st.markdown("### ⚙️ Search Settings")
    col1, col2 = st.columns(2)
    with col1:
        max_tokens = st.slider("🔤 Maximum response length", min_value=100, max_value=10000, value=2000)
        context_chunks = st.slider("📚 Context chunks to use", min_value=1, max_value=20, value=5, 
                                 help="More chunks provide deeper context but increase processing time")
    with col2:
        temperature = st.slider("🌡️ Response creativity", min_value=0.0, max_value=1.0, value=0.3, step=0.1)
        enable_citations = st.checkbox("📝 Include citations in response", value=True)
    
    # Common result display
    if search_type == "Website Content" and st.button("🔍 Search Website"):
        if not website_url or not website_query:
            st.warning("⚠️ Please enter both a website URL and a query.")
        else:
            with st.spinner(f"Analyzing website {website_url}..."):
                # Here you would call your function to process the website search
                search_prompt = f"""
                Analyze the website: {website_url} with crawl depth: {website_depth}.
                Question: {website_query}
                Respond with a detailed answer based on the website content.
                {"Include citations to specific pages where information was found." if enable_citations else ""}
                """
                search_result = generate_ai_content(search_prompt, st.session_state.api_key, st.session_state.api_model)
                
            st.success("✅ Search Complete!")
            st.text_area("🔍 Search Results:", search_result, height=300)
            
    elif search_type == "GitHub Repository" and st.button("🔍 Search Repository"):
        if not repo_url or not repo_query:
            st.warning("⚠️ Please enter both a repository URL and a query.")
        else:
            with st.spinner(f"Analyzing GitHub repository {repo_url}..."):
                # Here you would call your function to process the GitHub repo search
                included_content = ", ".join(include_options)
                search_prompt = f"""
                Analyze the GitHub repository: {repo_url}.
                Include: {included_content}
                File types to analyze: {file_types}
                Question: {repo_query}
                Respond with a detailed answer based on the repository content.
                {"Include citations to specific files where information was found." if enable_citations else ""}
                """
                search_result = generate_ai_content(search_prompt, st.session_state.api_key, st.session_state.api_model)
                
            st.success("✅ Search Complete!")
            st.text_area("🔍 Search Results:", search_result, height=300)
            
    elif search_type == "Deep Research" and st.button("🔍 Begin Deep Research"):
        if not deep_research_topic:
            st.warning("⚠️ Please enter a research topic.")
        else:
            with st.spinner(f"Conducting deep research on: {deep_research_topic}..."):
                # Process the deep research request
                search_prompt = f"""
                Deep research on {deep_research_topic}.
                please, think deeply and  provide a deep comprehensive analysis in at least more than 9000 words and add few emojis too with proper styling.
                {"Include citations to relevant sources where applicable." if enable_citations else ""}
                """
                search_result = generate_ai_content(search_prompt, st.session_state.api_key, st.session_state.api_model)
                st.session_state.search_result = search_result
                
            st.success("✅ Deep Research Complete!")
            st.text_area("🔍 Research Results:", search_result, height=500)
            
    elif search_type == "Document (PDF/TXT)" and st.button("🔍 Search Document"):
        if not uploaded_file or not doc_query:
            st.warning("⚠️ Please upload a document and enter a query.")
        else:
            with st.spinner(f"Analyzing document {uploaded_file.name}..."):
                # Here you would call your function to process the document search
                search_prompt = f"""
                Analyze the uploaded document: {uploaded_file.name}.
                Question: {doc_query}
                Respond with a detailed answer based on the document's content.
                {"Include citations to specific sections or pages where information was found." if enable_citations else ""}
                """
                search_result = generate_ai_content(search_prompt, st.session_state.api_key, st.session_state.api_model)
                
            st.success("✅ Search Complete!")
            st.markdown(f"<div style='height:300px; overflow-y:scroll; padding:10px; border:1px solid #ddd; border-radius:8px; background-color:#f9f9f9;'>{search_result}</div>", unsafe_allow_html=True)

    
    # Download & Copy Options for results
    if st.session_state.get("search_result"):
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("📥 Download Results", st.session_state.search_result, "search_results.txt")
        with col2:
            st.button("📋 Copy to Clipboard", on_click=lambda: st.write(
                "<script>navigator.clipboard.writeText(`" + st.session_state.search_result.replace("`", "\\`") + "`);</script>", 
                unsafe_allow_html=True))


# Advanced options expander
with st.expander("⚙️ Advanced Options"):
    # Document structure tabs
    tabs = st.tabs(["Template", "Content Style", "Structure", "Format", "Generation"])
    
    with tabs[0]:
        template_edit = st.text_area(
            "Customize prompt template (use {prompt} as placeholder for your input):",
            value=st.session_state.prompt_templates.get(selected_tool, "Create {prompt}"),
            height=100
        )
        st.session_state.prompt_templates[selected_tool] = template_edit
        
        # Template presets
        st.divider()
        col1, col2 = st.columns([3, 1])
        with col1:
            template_presets = {
                "Standard": "Create {prompt}",
                "Detailed Analysis": "Conduct a comprehensive analysis of {prompt}, including key insights, trends, and recommendations.",
                "Creative Writing": "Write a creative and engaging piece about {prompt} with vivid descriptions and compelling narrative.",
                "Business Report": "Generate a professional business report about {prompt} with executive summary, analysis, and actionable recommendations.",
                "Technical Guide": "Create a detailed technical guide for {prompt} with step-by-step instructions, code examples, and troubleshooting tips."
            }
            selected_preset = st.selectbox("Template Presets:", list(template_presets.keys()))
        with col2:
            if st.button("Apply Preset"):
                st.session_state.prompt_templates[selected_tool] = template_presets[selected_preset]
                st.rerun()

    
    with tabs[1]:
        # Writing style options
        col1, col2 = st.columns(2)
        with col1:
            response_length = st.selectbox(
                "Response Length:", ["Short", "Medium", "Long", "Very Short", "Concise", "Brief", "Detailed", "Extensive", "Summary", "In-depth", "Comprehensive", "Elaboration", "Thorough"]
,index=1  # Default to Medium
            )
            tone = st.selectbox(
                "Tone:",
                ["Professional", "Casual", "Academic", "Persuasive", "Inspirational", 
                 "Technical", "Conversational", "Humorous", "Formal", "Storytelling", 
                 "Authoritative", "Empathetic", "Neutral", "Enthusiastic", "Thoughtful"]
            )
            voice = st.selectbox(
                "Voice:",
                ["Active", "Passive", "First Person", "Second Person", "Third Person", 
                 "Objective", "Subjective", "Instructional", "Narrative", "Analytical"]
            )
            
        with col2:
            audience = st.selectbox(
                "Target Audience:",
                ["General", "Technical", "Executive", "Academic", "Marketing", 
                 "Education", "Healthcare", "Financial", "Legal", "Scientific",
                 "Children", "Teenagers", "Senior Management", "Beginners", "Advanced"]
            )
            industry = st.selectbox(
                "Industry Focus:",
                ["General", "Technology", "Healthcare", "Finance", "Education", 
                 "Entertainment", "Legal", "Marketing", "Science", "Engineering", 
                 "Manufacturing", "Retail", "Government", "Nonprofit", "Energy"]
            )
        
        # Language style 
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            language_style = st.selectbox(
                "Language Style:",
                ["Standard", "Simple", "Technical", "Poetic", "Journalistic", 
                 "Business", "Academic", "Conversational", "Persuasive", "Narrative"]
            )
            emotion = st.select_slider(
                "Emotional Intensity:",
                options=["Neutral", "Subtle", "Moderate", "Strong", "Intense"],
                value="Moderate"
            )
        with col2:
            formality = st.select_slider(
                "Formality Level:",
                options=["Very Casual", "Casual", "Neutral", "Formal", "Very Formal"],
                value="Neutral"
            )
            persona = st.selectbox(
                "Writing Persona:",
                ["Default", "Expert", "Teacher", "Coach", "Journalist", 
                 "Storyteller", "Analyst", "Researcher", "Consultant", "Mentor"]
            )
    
    with tabs[2]:
        col1, col2 = st.columns(2)
        with col1:
            max_words = st.number_input("Maximum Word Count:", min_value=50, max_value=10000, value=500, step=50)
            min_words = st.number_input("Minimum Word Count:", min_value=10, max_value=5000, value=100, step=50)
            complexity = st.select_slider(
                "Language Complexity:",
                options=["Elementary", "Middle School", "High School", "College", "Graduate", "Expert", "Technical"],
                value="College"
            )
        
        with col2:
            keyword_inclusion = st.text_area("Keywords to Include (comma separated):", height=68)
            forbidden_words = st.text_area("Words to Avoid (comma separated):", height=68)
            readability_target = st.slider("Readability Score Target:", 
                                          min_value=0, max_value=100, value=60, step=5, 
                                          help="Lower = more complex, Higher = simpler")
        
        # Structure options
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            structure_type = st.selectbox(
                "Content Structure:",
                ["Standard", "Problem-Solution", "Comparison", "Chronological", 
                 "Cause-Effect", "Spatial", "Process", "Thesis-Led", "Data-Led", 
                 "Argumentative", "Descriptive", "Exploratory", "Instructional"]
            )
            emphasis = st.selectbox(
                "Content Emphasis:",
                ["Balanced", "Data-Focused", "Process-Focused", "Outcome-Focused", 
                 "Analysis-Focused", "Solution-Focused", "Context-Focused"]
            )
        with col2:
            sections = st.multiselect(
                "Required Sections:",
                ["Introduction", "Background", "Methodology", "Results", "Discussion", 
                 "Conclusion", "Executive Summary", "Recommendations", "References", 
                 "Appendix", "FAQ", "Glossary", "Abstract", "Literature Review", 
                 "Case Studies", "Implementation", "Limitations", "Future Work"]
            )
            argument_style = st.selectbox(
                "Argument Style:",
                ["Balanced", "Steelmanning", "Devil's Advocate", "Persuasive", 
                 "Exploratory", "Socratic", "Analytical", "Comparative"]
            )
    
    with tabs[3]:
        col1, col2 = st.columns(2)
        with col1:
            output_format = st.selectbox(
                "Output Format:",
                ["Standard Text", "Markdown", "HTML", "JSON", "CSV", "Outline", 
                 "Bullet Points", "Q&A Format", "Table", "Script Format", 
                 "Newsletter", "Blog Post", "Academic Paper", "Business Report", 
                 "Technical Documentation", "Speech/Presentation"]
            )
            citation_style = st.selectbox(
                "Citation Style:",
                ["None", "APA", "MLA", "Chicago", "IEEE", "Harvard", "Vancouver", 
                 "AMA", "ASA", "Bluebook", "CSE", "ACS", "NLM"]
            )
        
        with col2:
            layout_style = st.selectbox(
                "Layout Style:",
                ["Standard", "Minimal", "Hierarchical", "Segmented", "Web-Optimized", 
                 "Print-Optimized", "Mobile-Optimized", "Presentation", "Technical"]
            )
            visual_elements = st.multiselect(
                "Visual Elements:",
                ["None", "Tables", "Lists", "Blockquotes", "Code Blocks", "Headings", 
                 "Sub-headings", "Bold Emphasis", "Italics", "Horizontal Rules", 
                 "Indentation"]
            )
            
        # Media and formatting
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            formatting_style = st.selectbox(
                "Formatting Style:",
                ["Standard", "Minimal", "Academic", "Web", "Print", "Journalistic", 
                 "Technical Document", "Creative", "Business", "Scientific"]
            )
        with col2:
            syntax_highlighting = st.selectbox(
                "Code Syntax Highlighting:",
                ["None", "Standard", "GitHub", "VSCode", "Atom", "Sublime"]
            )
    
    with tabs[4]:
        # Generation parameters
        col1, col2 = st.columns(2)
        with col1:
            output_length = st.select_slider(
                "Output Detail Level:",
                options=["Minimal", "Brief", "Standard", "Detailed", "Comprehensive", "Exhaustive"],
                value="Standard"
            )
            creativity = st.slider("Creativity:", min_value=0.0, max_value=1.0, value=0.7, step=0.1,
                                help="Lower = more deterministic, Higher = more creative")
        
        with col2:
            determinism = st.slider("Determinism:", min_value=0.0, max_value=1.0, value=0.3, step=0.1,
                                help="Lower = more varied outputs, Higher = more consistent outputs")
            reasoning_depth = st.select_slider(
                "Reasoning Depth:",
                options=["Basic", "Standard", "Advanced", "Expert", "Comprehensive"],
                value="Standard",
                help="Controls depth of logical analysis and reasoning in generated content"
            )
        
        # Advanced AI parameters
        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            factuality = st.slider("Factuality Weight:", min_value=0.0, max_value=1.0, value=0.8, step=0.1,
                                help="Higher values prioritize factual accuracy over creativity")
            expertise_level = st.select_slider(
                "AI Expertise Level:",
                options=["Generalist", "Field Expert", "Subject Specialist", "Domain Authority", "World-Class Expert"],
                value="Field Expert",
                help="Level of expertise the AI should demonstrate in the response"
            )
        
        with col2:
            context_relevance = st.slider("Context Relevance:", min_value=0.0, max_value=1.0, value=0.9, step=0.1,
                                      help="Higher values keep content more focused on the specific prompt")
            innovation_level = st.select_slider(
                "Innovation Level:",
                options=["Conservative", "Balanced", "Progressive", "Cutting-Edge", "Revolutionary"],
                value="Balanced",
                help="Controls how novel or traditional the ideas and approach should be"
            )
        
        # Specialized content options
        st.divider()
        st.subheader("Specialized Content Options")
        st.info("Select additional content features to include in generation")
        col1, col2, col3 = st.columns(3)
        with col1:
            tech_analysis = st.checkbox("Enable Technical Analysis", help="Add technical depth with specialized terminology")
            data_viz = st.checkbox("Include Data Visualization", help="Add descriptions of charts/graphs where appropriate")
            case_studies = st.checkbox("Add Case Studies", help="Include relevant examples and case studies")
        
        with col2:
            research_focus = st.checkbox("Research Focus", help="Emphasize research findings and methodologies")
            step_breakdown = st.checkbox("Step-by-Step Breakdown", help="Include detailed procedural explanations")
            comparative = st.checkbox("Comparative Analysis", help="Include comparisons to alternatives or competitors")
        
        with col3:
            future_trends = st.checkbox("Future Trends", help="Include predictions and future developments")
            historical = st.checkbox("Historical Context", help="Add historical background and evolution")
            expert_citations = st.checkbox("Expert Citations", help="Include references to subject matter experts")

# Add style instructions to prompt based on selections
style_instructions = {
    "Response Length": response_length,
    "tone": tone,
    "voice": voice,
    "audience": audience,
    "industry": industry,
    "max_words": max_words,
    "min_words": min_words,
    "output_format": output_format,
    "structure_type": structure_type,
    "sections": sections,
    "output_length": output_length,
    "creativity": creativity,
    "reasoning_depth": reasoning_depth
}

# Store parameters to session state for use in prompt building
st.session_state.style_instructions = style_instructions
# Export/Import functionality
st.sidebar.markdown("---")
st.sidebar.markdown("### 💾 Export/Import History")

# Export history
if st.session_state.history and st.sidebar.button("📤 Export History"):
    history_json = json.dumps(st.session_state.history)
    b64_history = base64.b64encode(history_json.encode()).decode()
    href = f'<a href="data:application/json;base64,{b64_history}" download="ai_content_history.json">Download History File</a>'
    st.sidebar.markdown(href, unsafe_allow_html=True)

# Import history
uploaded_file = st.sidebar.file_uploader("Import History:", type=['json'])
if uploaded_file is not None:
    try:
        imported_history = json.loads(uploaded_file.read())
        if st.sidebar.button("📥 Load Imported History"):
            st.session_state.history = imported_history
            st.sidebar.success("History imported successfully!")
    except Exception as e:
        st.sidebar.error(f"Error importing history: {e}")

# Generate button and output area


if st.button("🚀 Generate Content", type="primary"):
    if not st.session_state.api_key:
        st.error("Please enter your API key in the sidebar first.")
    elif not user_prompt:
        st.warning("Please enter a prompt to generate content.")
    else:
        # Build prompt with template and style instructions
        template = st.session_state.prompt_templates.get(selected_tool, "Create {prompt}")
        formatted_prompt = template.replace("{prompt}", user_prompt)
        
        # Add style instructions to prompt
        style_prompt = f"{formatted_prompt}\n\nStyle Guidelines:\n"
        for key, value in st.session_state.style_instructions.items():
            if value and value not in ["None", "Standard"] and (not isinstance(value, list) or len(value) > 0):
                style_prompt += f"- {key.replace('_', ' ').title()}: {value}\n"
        
        # Generate content
        output = generate_ai_content(style_prompt, st.session_state.api_key, st.session_state.api_model)
        
        # Display output
        st.markdown("### 🎯 Generated Content")
        with st.expander("Content Output", expanded=True):
            st.markdown(output)
        
        # Save to history
        save_to_history(selected_tool, user_prompt, output)

        # Export options
        st.markdown("### 📥 Export Options")

        # Function to create DOCX file
        def create_docx(text):
            doc = docx.Document()
            doc.add_paragraph(text)
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)  # Move to the beginning of the file
            return doc_stream

        # Function to create PDF file
        import fitz  # PyMuPDF for PDF generation
        def create_pdf(text):
            pdf_stream = BytesIO()
            doc = fitz.open()  # Create a new PDF
            page = doc.new_page(width=595, height=842)  # A4 size
            text = text.replace("\n", "<br>")  # Maintain line breaks
            page.insert_text((50, 50), text, fontsize=12, color=(0, 0, 0))
            doc.save(pdf_stream)
            pdf_stream.seek(0)
            return pdf_stream

        # Download buttons
        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button("📄 Download as TXT", output, "generated_content.txt")
        with col2:
            docx_file = create_docx(output)
            st.download_button("📄 Download as DOCX", docx_file, "generated_content.docx")
        with col3:
            pdf_file = create_pdf(output)
            st.download_button("📄 Download as PDF", pdf_file, "generated_content.pdf")



              
# Add theme selector
st.sidebar.markdown("---")
st.sidebar.markdown("### 🎨 App Theme")
themes = ["Light", "Dark", "Blue", "Green", "Purple"]
selected_theme = st.sidebar.selectbox("Select Theme:", themes, index=0)

# Apply selected theme with custom CSS
theme_colors = {
    "Light": {"bg": "#f8f9fa", "accent": "#3b82f6"},
    "Dark": {"bg": "#1e293b", "accent": "#8b5cf6"},
    "Blue": {"bg": "#f0f9ff", "accent": "#0284c7"},
    "Green": {"bg": "#f0fdf4", "accent": "#16a34a"},
    "Purple": {"bg": "#faf5ff", "accent": "#9333ea"}
}

theme_css = f"""
<style>
    .main {{ background-color: {theme_colors[selected_theme]["bg"]} }}
    .stButton>button {{ background: linear-gradient(90deg, {theme_colors[selected_theme]["accent"]}, {theme_colors[selected_theme]["accent"]}88) }}
</style>
"""
st.sidebar.markdown(theme_css, unsafe_allow_html=True)
#g
